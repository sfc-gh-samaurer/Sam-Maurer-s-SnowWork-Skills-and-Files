#!/usr/bin/env python3
"""
Snowflake PS Project Attachment Generator

Generates the two standard project-specific attachment documents:
  1. Data Migration & Validation Project (DMVA)
  2. Code Conversion Project

User-supplied values are wrapped in OOXML <w:ins> (tracked insertions) so that
the PS writer can review and accept/reject changes in Word before sending to
the customer.

Usage:
    python generate_project_attachments.py <json_input_path> <output_dir>

The JSON must contain one or both of:
  - "dmva_attachment": { ... }
  - "code_conversion_attachment": { ... }

See SKILL.md for full JSON schemas.
"""

import json
import sys
import os
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from lxml import etree

# ── Constants ─────────────────────────────────────────────────────────────────
FONT_NAME   = "Arial"
FONT_SIZE   = Pt(7.5)
FONT_SZ_HP  = "15"          # half-points for <w:sz>
PAGE_WIDTH  = 7772400
PAGE_HEIGHT = 10058400
MARGIN_LR   = 457200
MARGIN_TOP  = 337185
MARGIN_BOTTOM = 640080
TABLE_WIDTH_DXA = 10800
LINE_SPACING = 170
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

TRACK_AUTHOR = "Snowflake PS Generator"
TRACK_DATE   = f"{date.today().isoformat()}T00:00:00Z"

_ins_id_counter = 0


def _next_ins_id():
    global _ins_id_counter
    _ins_id_counter += 1
    return _ins_id_counter


# ── Document helpers ───────────────────────────────────────────────────────────

def create_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width  = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.left_margin = MARGIN_LR
    section.right_margin = MARGIN_LR
    section.top_margin  = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    if doc.paragraphs:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)
    return doc


def _rPr_xml(bold=False):
    b = "<w:b/><w:bCs/>" if bold else ""
    return (
        f'<w:rPr xmlns:w="{W_NS}">'
        f'  <w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:eastAsia="{FONT_NAME}" w:cs="{FONT_NAME}"/>'
        f'  <w:sz w:val="{FONT_SZ_HP}"/><w:szCs w:val="{FONT_SZ_HP}"/>'
        f'  {b}'
        f'</w:rPr>'
    )


def _apply_para_spacing(para):
    pPr = para._p.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        pPr.append(parse_xml(
            f'<w:spacing {nsdecls("w")} w:before="0" w:after="0" '
            f'w:line="{LINE_SPACING}" w:lineRule="auto"/>'
        ))
    else:
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'),  '0')
        sp.set(qn('w:line'),   str(LINE_SPACING))
        sp.set(qn('w:lineRule'), 'auto')


def _add_plain_run(para, text, bold=False):
    """Add a normal (un-tracked) run."""
    run = para.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rPr.insert(0, parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" '
            f'w:eastAsia="{FONT_NAME}" w:cs="{FONT_NAME}"/>'
        ))
    return run


def _add_tracked_run(para, text, bold=False):
    """Add a tracked insertion run (<w:ins>) to an existing paragraph."""
    ins_id = _next_ins_id()
    ins_xml = (
        f'<w:ins xmlns:w="{W_NS}" w:id="{ins_id}" '
        f'w:author="{TRACK_AUTHOR}" w:date="{TRACK_DATE}">'
        f'  <w:r>'
        f'    {_rPr_xml(bold)}'
        f'    <w:t xml:space="preserve">{_escape(text)}</w:t>'
        f'  </w:r>'
        f'</w:ins>'
    )
    para._p.append(parse_xml(ins_xml))


def _escape(text):
    return (str(text)
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;'))


def add_para(doc, text="", bold=False, tracked=False):
    para = doc.add_paragraph()
    _apply_para_spacing(para)
    if text:
        if tracked:
            _add_tracked_run(para, text, bold=bold)
        else:
            _add_plain_run(para, text, bold=bold)
    return para


def add_bullet(doc, text, tracked=False):
    para = doc.add_paragraph()
    _apply_para_spacing(para)
    bullet = "\u2022 "
    _add_plain_run(para, bullet)
    if tracked:
        _add_tracked_run(para, text)
    else:
        _add_plain_run(para, text)
    return para


def add_blank(doc):
    return add_para(doc)


def add_heading(doc, text, level=1):
    return add_para(doc, text, bold=(level == 1))


# ── Table helpers ──────────────────────────────────────────────────────────────

def _set_table_style(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    # Width
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblPr.append(parse_xml(
            f'<w:tblW {nsdecls("w")} w:w="{TABLE_WIDTH_DXA}" w:type="dxa"/>'
        ))
    else:
        tblW.set(qn('w:w'), str(TABLE_WIDTH_DXA))

    # Borders
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top    w:val="single" w:sz="4" w:color="000000"/>'
        '  <w:left   w:val="single" w:sz="4" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:color="000000"/>'
        '  <w:right  w:val="single" w:sz="4" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:color="000000"/>'
        '</w:tblBorders>'
    ))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def _set_col_widths(table, widths):
    if not widths:
        return
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(widths):
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                    tc.insert(0, tcPr)
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcPr.append(parse_xml(
                        f'<w:tcW {nsdecls("w")} w:w="{widths[ci]}" w:type="dxa"/>'
                    ))
                else:
                    tcW.set(qn('w:w'), str(widths[ci]))
                    tcW.set(qn('w:type'), 'dxa')


def _fill_cell(cell, text, bold=False, tracked=False):
    """Write text into a table cell, clearing existing content."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    _apply_para_spacing(para)
    if tracked:
        _add_tracked_run(para, str(text), bold=bold)
    else:
        _add_plain_run(para, str(text), bold=bold)


def add_table(doc, headers, rows, col_widths=None, header_bold=True, tracked_cols=None):
    """
    tracked_cols: set of column indices whose data cells should be tracked insertions.
    Header cells are never tracked.
    """
    tracked_cols = tracked_cols or set()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    _set_table_style(table)

    for ci, h in enumerate(headers):
        _fill_cell(table.rows[0].cells[ci], h, bold=header_bold)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            _fill_cell(
                table.rows[ri + 1].cells[ci],
                val,
                tracked=(ci in tracked_cols)
            )

    if col_widths:
        _set_col_widths(table, col_widths)

    return table


# ═══════════════════════════════════════════════════════════════════════════════
# DATA MIGRATION ATTACHMENT
# ═══════════════════════════════════════════════════════════════════════════════

def generate_dmva_attachment(data: dict, output_path: str) -> str:
    """
    Generate the Data Migration & Validation Project attachment.

    Expected keys in data:
      legacy_platform       str   e.g. "Amazon Redshift (3 clusters, 19 databases)"
      target_platform       str   default "Snowflake Service (non-production databases)"
      engagement_type       str   "t_and_m" | "fixed_fee"
      scope:
        tables_full_load    str/int
        tables_incremental  str/int
        total_data_volume   str   e.g. "4.2 TB"
        oversized_columns   str/int
        modification_columns str/int
        notes               str   (optional)
      resource_plan: list of {week, activity, sf_hours, customer_hours}
      extra_exclusions:     list[str]  (optional)
      extra_assumptions:    list[str]  (optional)
    """
    global _ins_id_counter
    _ins_id_counter = 0

    doc = create_document()

    legacy  = data.get("legacy_platform", "[Legacy Platform]")
    target  = data.get("target_platform", "Snowflake Service (non-production databases)")
    eng     = data.get("engagement_type", "t_and_m")
    scope   = data.get("scope", {})

    # ── Title ─────────────────────────────────────────────────────────────────
    add_heading(doc, "Attachment 1: Data Migration & Validation Project", level=1)
    add_blank(doc)

    # ── Opening ───────────────────────────────────────────────────────────────
    add_para(doc, "Data Migration Services:")
    add_blank(doc)
    add_para(doc,
        "During the Term, Snowflake will perform the Technical Services below related to a "
        "one-time data migration from the Legacy Platform(s) listed below to the Snowflake "
        "Service (\"Data Migration Services\")."
    )
    add_blank(doc)

    # ── i. Scope ──────────────────────────────────────────────────────────────
    add_para(doc, "i.  Data Migration Scope.", bold=True)
    add_para(doc,
        "The following table outlines the scope of data to be migrated from the Legacy "
        "Platform to the Target Platform, subject to the availability of hours in this SOW."
    )
    add_blank(doc)
    add_para(doc, "Table 1: Data Migration Scope Parameters", bold=True)
    add_blank(doc)

    scope_rows = [
        ["Legacy Platform",                      legacy],
        ["Target Platform",                      target],
        ["# of Tables (Full Load Process)",      str(scope.get("tables_full_load", "XX"))],
        ["# of Tables (Incremental Load Process)", str(scope.get("tables_incremental", "XX"))],
        ["Total Data Volume (TB)",               str(scope.get("total_data_volume", "XX"))],
        ["Oversized Values (# of Columns)",      str(scope.get("oversized_columns", "XX"))],
        ["Modification of Data (# of Columns)",  str(scope.get("modification_columns", "XX"))],
    ]
    if scope.get("notes"):
        scope_rows.append(["Notes", scope.get("notes")])

    add_table(
        doc,
        ["Scope Parameter", "Value"],
        scope_rows,
        col_widths=[5400, 5400],
        header_bold=True,
        tracked_cols={1},   # value column is tracked
    )
    add_blank(doc)

    # ── ii. Snowflake Responsibilities ────────────────────────────────────────
    add_para(doc, "ii.  Snowflake Responsibilities:", bold=True)
    add_para(doc,
        "Snowflake will perform the following data-migration related tasks subject to the "
        "availability of hours in the Order Form, if any:"
    )
    add_blank(doc)

    sf_responsibilities = [
        ("Create Data Migration Execution Plan",
         "With the help of the Customer team, Snowflake will create a detailed data migration "
         "plan showing the tables to be migrated per project phase. The plan will also describe "
         "the data validation timelines."),
        ("Data Migration Definition",
         "Snowflake will work with the Customer team to define the following for the data "
         "migration project: Requirements for Modification of Data, Incremental Table Design "
         "Criteria, and Validation Criteria (including row counts, distinct value counts, "
         "min/max values, and column sum values for numeric columns)."),
        ("Data Mapping",
         "Provide the Customer with data mapping scripts to run against all columns in their "
         "production Legacy Platform system. Any column containing large LOB data which will "
         "not fit in a Varchar or Binary in Snowflake is considered an Oversized Value Column "
         "and is not in scope for automated data migration or validation."),
        ("Data Migration",
         "Snowflake will only be responsible for the one-time migration of data and validation "
         "of the migration process and migrated data from a non-production environment to the "
         "Target Platform. Snowflake will perform all activities for this data migration project "
         "in non-production environments."),
        ("Data Validation",
         "Perform data validation integrity checks between the Legacy Platform and the Target "
         "Platform post data migration and provide a data quality report for sign-off."),
        ("Progress Reporting",
         "Snowflake delivery manager will provide a weekly update of progress against the agreed "
         "to data migration execution plan."),
    ]

    for title, body in sf_responsibilities:
        add_para(doc, f"{title}:  {body}")
        add_blank(doc)

    # ── iii. Customer Responsibilities ────────────────────────────────────────
    add_para(doc, "iii.  Customer Responsibilities:", bold=True)
    add_para(doc,
        "Snowflake's performance of Data Migration Services is dependent on Customer's timely "
        "performing of the following activities:"
    )
    add_blank(doc)
    add_para(doc, "Customer is solely responsible for the following tasks:")

    customer_responsibilities = [
        "Notifying Snowflake in advance of any changes to the table structures or substantial changes in average data volume being migrated as part of this SOW.",
        "Ensuring that no medical data, PHI, PII or other similarly sensitive data is accessible to Snowflake or within the data migration scope for Snowflake.",
        "Assist with creation of the data migration execution plan by providing information on table load dependencies, data loading schedules, business priorities, etc.",
        "Provide any requirements to modify data as part of the data migration.",
        "Provide validation integrity check requirements for Snowflake to review and confirm agreement.",
        "Run data profiling scripts against the Legacy Platform data and provide results to Snowflake.",
        "Provide information on table design for any tables requiring incremental load process such as primary key definitions or columns used to identify recently updated records.",
        "Provide read access accounts to the non-production Legacy Platform for purposes of data extraction and data validation.",
        "Provide read/write access accounts to non-production Target Platform for purposes of data loading and data validation.",
        "Provide a virtual desktop environment or customer laptop from which to execute all data migration and data validation tasks.",
        "Provide the configuration information for where extracted data files will be staged in the data migration process.",
        "Provide instructions on Legacy Platform and Target Platform locations for data storage, and review and approve all extraction scripts to ensure data is migrated to appropriate locations.",
        "Assure adequate system bandwidth to perform data extraction and to run data validation processes on the Legacy Platform.",
        "Assure adequate network bandwidth to transfer data extract files to the Target Platform.",
        "Ensure that any amount of test data (non-production data) for data extraction of a table is representative of the production environment.",
        "Execution of the data migration and data validation activities into the Customer non-production Snowflake environment.",
        "Providing a set of fully static data during the development, configuration and testing of the data migration process.",
        "Providing a suitable extraction window for all tables where data will be static when executing the full data migration and performing the data validation.",
        "Providing configuration specifications for the migration and reviewing the migration plan, scripts and configuration files to ensure it meets internal or external requirements.",
        "Customer is solely responsible and accountable for decisions related to the amount of validation testing performed.",
        "Sign-off of the data validation report.",
        "Any further validation or testing required for Customer's particular use or application of the migrated data.",
    ]

    for item in customer_responsibilities:
        add_bullet(doc, item)
    add_blank(doc)

    # ── iv. Resource Plan ─────────────────────────────────────────────────────
    add_para(doc, "iv.  Data Migration Resource Plan:", bold=True)
    add_blank(doc)

    resource_plan = data.get("resource_plan", [])
    if resource_plan:
        headers = ["Week", "Activity", "Snowflake Hours", "Customer Hours"]
        rows = [
            [
                str(r.get("week", "")),
                str(r.get("activity", "")),
                str(r.get("sf_hours", "")),
                str(r.get("customer_hours", "")),
            ]
            for r in resource_plan
        ]
        add_table(doc, headers, rows, col_widths=[1200, 5400, 2100, 2100],
                  tracked_cols={0, 1, 2, 3})
    else:
        add_para(doc, "[Resource plan by week to be inserted]", tracked=True)
    add_blank(doc)

    # ── v. Exclusions ─────────────────────────────────────────────────────────
    add_para(doc, "v.  Technical Scope Exclusions:", bold=True)
    add_blank(doc)

    standard_exclusions = [
        "Physical transfer of any data.",
        "Non-scripted data validation.",
        "Oversized Value Column.",
        "Validation of views, functions, procedures or any other objects other than tables.",
        "Validation of end-user reporting.",
        "Analyzing or optimizing network bandwidth.",
        "Migration of data to the Customer Snowflake Production Environment.",
        "Unstructured data.",
        "Referenced data that is outside of the Legacy Platform.",
    ]

    for item in standard_exclusions:
        add_bullet(doc, item)

    for item in data.get("extra_exclusions", []):
        add_bullet(doc, item, tracked=True)
    add_blank(doc)

    # ── vi. Assumptions ───────────────────────────────────────────────────────
    add_para(doc, "vi.  Technical Scope Assumptions:", bold=True)
    add_blank(doc)

    standard_assumptions = [
        "Environment variables required for DMVA are stored on the same Virtual Machine as DMVA. Use and configuration of any vault or secrets management solution is out of scope.",
    ]

    for item in standard_assumptions:
        add_bullet(doc, item)

    for item in data.get("extra_assumptions", []):
        add_bullet(doc, item, tracked=True)
    add_blank(doc)

    # ── vii. Service Type ─────────────────────────────────────────────────────
    add_para(doc, "vii.  Service Type:", bold=True)
    add_blank(doc)

    if eng == "fixed_fee":
        add_para(doc,
            "The Data Migration Services will be provided on a fixed fee basis. Migration "
            "activities will not exceed the limits and object counts specified in Table 1 above."
        )
    else:
        add_para(doc,
            "The Data Migration Services are provided on a time and material basis and are "
            "subject to the availability of hours for the Data Migration Services as specified "
            "in the Order Form."
        )
    add_blank(doc)

    doc.save(output_path)
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
# CODE CONVERSION ATTACHMENT
# ═══════════════════════════════════════════════════════════════════════════════

def generate_code_conv_attachment(data: dict, output_path: str) -> str:
    """
    Generate the Code Conversion Project attachment.

    Expected keys in data:
      legacy_platform           str   e.g. "Amazon Redshift"
      assessment_date           str   e.g. "March 2026" | "Not Performed"
      target_platform           str   default "Snowflake Service (non-production databases)"
      engagement_type           str   "t_and_m" | "fixed_fee"
      objects: list of {type, count}
        Standard types: Tables, Views, Materialized Views, Macros, Functions,
                        Procedures, Packages, Triggers, SQL Scripts, Utility Scripts
      refactoring: list of {type, estimated, max_refactor}
        Use standard type names from Table 2; add extras as needed
      prerequisites:
        snowconvert_required    bool  default True
      unit_test:
        outside_env_dummy       bool
        without_data_customer   bool
        dummy_data_customer     bool
        customer_static_data    bool
      code_correction_support_weeks   int   (T&M: use 0 for "end of engagement")
      staffing_plan: list of {week, activity, sf_hours}
      extra_exclusions: list[str]
    """
    global _ins_id_counter
    _ins_id_counter = 0

    doc = create_document()

    legacy   = data.get("legacy_platform", "[Legacy Platform]")
    assess   = data.get("assessment_date", "Insert Date / Not Performed")
    target   = data.get("target_platform", "Snowflake Service (non-production databases)")
    eng      = data.get("engagement_type", "t_and_m")
    prereqs  = data.get("prerequisites", {})
    unit_test = data.get("unit_test", {})

    # ── Title ─────────────────────────────────────────────────────────────────
    add_heading(doc, "Attachment 1: Code Conversion Project", level=1)
    add_blank(doc)

    # ── Opening ───────────────────────────────────────────────────────────────
    add_para(doc, "Code Conversion Services:")
    add_blank(doc)
    add_para(doc,
        "During the Term, Snowflake will perform the Technical Services below related to code "
        "conversion from the Legacy Platform(s) listed below to the Snowflake Service "
        "(\"Code Conversion Services\")."
    )
    add_blank(doc)

    # ── i. Object and Script Conversion Scope ─────────────────────────────────
    add_para(doc, "i.  Object and Script Conversion Scope.", bold=True)
    add_para(doc,
        "The following database objects (\"Objects\") from the legacy platform(s) and "
        "environment(s) listed in the table below are in-scope for the Code Conversion "
        "Services, subject to the availability of hours in the Order Form."
    )
    add_blank(doc)
    add_para(doc, "Table 1: Objects and Legacy Platforms", bold=True)
    add_blank(doc)

    # Build header rows (platform info)
    platform_rows = [
        ["Legacy Platform", legacy],
        ["Code Assessment Completed On", assess],
        ["Target Platform", target],
    ]

    # Object type rows
    standard_obj_types = [
        "Tables", "Views", "Materialized Views", "Macros", "Functions",
        "Procedures", "Packages", "Triggers",
        "SQL Scripts \u2013 Wrapped/Unwrapped",
        "Teradata Utility Scripts \u2013 Wrapped/Unwrapped",
        "Oracle PL/SQL Scripts \u2013 Wrapped/Unwrapped",
    ]

    user_objects = {o.get("type", ""): o.get("count", "0") for o in data.get("objects", [])}

    col_label = "Object Count" if eng == "t_and_m" else "Not to Exceed Object Count"
    object_rows = []
    for ot in standard_obj_types:
        cnt = user_objects.get(ot, "0")
        object_rows.append([ot, str(cnt)])

    # Any extra object types not in the standard list
    for obj in data.get("objects", []):
        if obj.get("type") not in standard_obj_types:
            object_rows.append([obj.get("type", ""), str(obj.get("count", "0"))])

    # Combine into one table: header rows + object rows
    all_rows = platform_rows + [["Object Type", col_label]] + object_rows
    tbl = doc.add_table(rows=len(all_rows), cols=2)
    _set_table_style(tbl)
    _set_col_widths(tbl, [5400, 5400])

    for ri, row_data in enumerate(all_rows):
        for ci, val in enumerate(row_data):
            cell = tbl.rows[ri].cells[ci]
            for p in cell.paragraphs:
                for r in p.runs:
                    r.text = ""
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            _apply_para_spacing(para)

            is_section_header = (row_data in [["Object Type", col_label]] or
                                  row_data[0] in ["Legacy Platform", "Code Assessment Completed On", "Target Platform"])
            is_value_col = (ci == 1)
            is_object_row = (ri >= len(platform_rows) + 1)

            if is_section_header:
                _add_plain_run(para, val, bold=True)
            elif is_value_col and is_object_row:
                _add_tracked_run(para, val)
            elif ci == 1 and ri < len(platform_rows):
                _add_tracked_run(para, val)
            else:
                _add_plain_run(para, val)

    add_blank(doc)

    # ── ii. Snowflake Responsibilities ────────────────────────────────────────
    add_para(doc, "ii.  Snowflake Responsibilities:", bold=True)
    add_para(doc,
        "For the Objects identified above, and subject to the exclusions and Customer "
        "responsibilities below, Snowflake will, subject to the availability of hours, "
        "perform the following activities as part of Code Conversion Services:"
    )
    add_blank(doc)

    sf_responsibilities = [
        ("Code Conversion Assessment",
         "Snowflake will perform an automated assessment for the Customer code base to provide "
         "Customer with an understanding of the conversion complexities and aid the Customer "
         "and Snowflake in the creation of a detailed conversion plan."),
        ("Create Detailed Conversion Plan",
         "With the help of the Customer team, Snowflake will create a detailed conversion plan "
         "breaking down the Objects per project phase. The plan will also describe the testing "
         "and UAT timelines where applicable."),
        ("Automated Code Conversion",
         "Snowflake will run an automated code conversion process on Objects and review "
         "resulting converted Objects."),
        ("Conversion to Snowflake Scripting",
         "All macros, functions, and procedures will be converted to Snowflake Service "
         "scripting (as opposed to Snowflake Javascript) when functionally viable, based upon "
         "existing and available capabilities of Snowflake Service scripting."),
        ("Create Dummy Test Data",
         "If Customer is not providing real test data for Snowflake to unit test objects, "
         "Snowflake will generate a set of dummy data to use for unit testing. Customer "
         "acknowledges that dummy data generated by Snowflake may not meet all test conditions."),
        ("Compile Converted Objects",
         "Snowflake will compile converted Objects within a non-Customer Snowflake Service "
         "database or system or within a Customer non-production environment if access has been "
         "given to Snowflake resources."),
    ]

    for title, body in sf_responsibilities:
        add_para(doc, f"{title}:  {body}")
        add_blank(doc)

    # Code Refactoring
    add_para(doc,
        "Code Refactoring:  Refactoring of code may be needed when converting code statements "
        "like-for-like to Snowflake SQL is not possible or may result in poor code performance. "
        "The following table shows the number and types of code objects identified that may "
        "require refactoring and the maximum number that Snowflake will attempt to refactor as "
        "part of the Code Conversion Services."
    )
    add_blank(doc)
    add_para(doc, "Table 2: Code Refactoring", bold=True)
    add_blank(doc)

    standard_refactor_types = [
        "Case Insensitive Collation",
        "Cursor Loops",
        "Global Temporary Tables",
        "Materialized Views",
        "Nested Functions",
        "Nested Procedures and Transactions",
        "Non-Logic View Layers",
        "Object Generating Dynamic SQL",
        "Triggers",
        "Renaming of Objects",
        "Custom Refactor Request",
    ]

    user_refactor = {
        r.get("type", ""): (str(r.get("estimated", "0")), str(r.get("max_refactor", "0")))
        for r in data.get("refactoring", [])
    }

    max_exceed_label = "Maximum Instances to be Refactored by Snowflake" if eng == "t_and_m" else \
                       "Maximum Instances to be Refactored by Snowflake (Not to Exceed)"
    refactor_rows = []
    for rt in standard_refactor_types:
        est, mx = user_refactor.get(rt, ("0", "0"))
        refactor_rows.append([rt, str(est), str(mx)])

    for r in data.get("refactoring", []):
        if r.get("type") not in standard_refactor_types:
            refactor_rows.append([r.get("type", ""), str(r.get("estimated", "0")), str(r.get("max_refactor", "0"))])

    add_table(
        doc,
        ["Type", "Estimated Number of Instances Identified", max_exceed_label],
        refactor_rows,
        col_widths=[4200, 3000, 3600],
        header_bold=True,
        tracked_cols={1, 2},
    )
    add_blank(doc)
    add_para(doc, "Any of the Code Refactoring line items that are not applicable (zero quantity) should have a 0 quantity.")
    add_blank(doc)

    # Functional Unit Test
    add_para(doc, "Functional Unit Test:  Snowflake will conduct unit test of converted Objects as follows:")
    add_blank(doc)

    def _yn(val):
        return "yes" if val else "no"

    ut_options = [
        (_yn(unit_test.get("outside_env_dummy",      False)), "Test outside of Customer's development environment with dummy data"),
        (_yn(unit_test.get("without_data_customer",  False)), "Test without data in Customer's development environment"),
        (_yn(unit_test.get("dummy_data_customer",    False)), "Test with dummy data in Customer's development environment"),
        (_yn(unit_test.get("customer_static_data",   True)),  "Test with Customer-provided static data in Customer's development environment. Does not include data validation or data matching with source system."),
    ]

    for yn_val, description in ut_options:
        para = doc.add_paragraph()
        _apply_para_spacing(para)
        _add_tracked_run(para, f" {yn_val} ", bold=True)
        _add_plain_run(para, f" / {'yes' if yn_val == 'no' else 'no'}  \u2013  {description}")

    add_blank(doc)
    add_para(doc, "Note: dummy data does not guarantee that all test conditions will be met.")
    add_blank(doc)

    add_para(doc,
        "Delivery of Converted Objects:  Snowflake will deliver a set of executable converted "
        "objects along with deployment instructions to Customer."
    )
    add_blank(doc)

    # ── iii. Customer Responsibilities ────────────────────────────────────────
    add_para(doc, "iii.  Customer Responsibilities:", bold=True)
    add_para(doc,
        "Snowflake's performance of Code Conversion Services is dependent on Customer timely "
        "performing the following activities:"
    )
    add_blank(doc)

    # Prerequisites table
    add_para(doc, "Prerequisites to start Code Conversion:", bold=True)
    add_blank(doc)

    snowconvert_req = "Yes" if prereqs.get("snowconvert_required", True) else "No"

    prereq_rows = [
        ["Install SnowConvert (\"Conversion Software\") subject to separate terms and conditions",
         snowconvert_req],
        ["Setup Snowflake Service Environment \u2013 Database, Roles, Apply Masking Policies on sensitive columns", "Yes"],
        ["Ensure latest code inventory is provided to Snowflake team", "Yes"],
        ["Provide inventory object list (template provided by Snowflake). Include any prioritization requirements.", "Yes"],
        ["Onboard Snowflake resources and provide all required access to Customer's systems", "Yes"],
        ["Provide end to end migration project plan", "Yes"],
        ["Code Conversion Project Management: Setup issue tracking tool, setup code repository", "Yes"],
        ["Assemble team for code conversion, data migration, QA activities. Provide list of team members.", "Yes"],
        ["POC and Tool Selection \u2013 Perform POCs to build/choose data migration, data validation tools", "Optional"],
        ["Finalize QA Testing Approach \u2013 Select data validation tool, success criteria & validation scripts", "Optional"],
        ["Perform one-time data migration (for code conversion unit testing with Customer data)", "Optional"],
        ["To assist testing, if UDFs are in scope for conversion, provide the base tables and data used", "Yes"],
    ]

    add_table(
        doc,
        ["Customer Prerequisite", "Required"],
        prereq_rows,
        col_widths=[8600, 2200],
        header_bold=True,
        tracked_cols={1},
    )
    add_blank(doc)

    customer_responsibilities_cc = [
        "Code Collection: Gathering of the code Objects required to be translated. Customer will execute scripts provided by Snowflake to extract all Objects.",
        "Remediation of Missing Objects: Customer is responsible for providing a complete code base. Missing code may result in extra effort. Customer must resolve any missing dependent objects within one (1) day of notification by Snowflake.",
        "Loading of Unit Test Data: Customer is responsible for loading test data into the converted Snowflake tables provided by Snowflake prior to starting unit testing.",
        "Deploy and Execute Converted Objects: Deploy and execute converted Objects in Customer's Snowflake Service account and/or other applicable system.",
        "Remediation of Functional Gaps: Code Conversion Services do not cover certain code scenarios and functionality. Components requiring architectural redesign are the Customer's responsibility.",
        "Data Validation / System Integration Testing: Customer is responsible and accountable for system integration testing and data validation of all converted Objects.",
        "Pre-Production Testing: Customer is responsible for any user acceptance testing, operational readiness testing and other testing required for production release.",
        "Project/Program Management: Customer is responsible for providing overall project management for necessary activities related to the Customer's migration.",
        "Renaming: If renaming of Objects is part of the scope, Customer must provide all renaming requirements in the first 2 weeks of the project.",
    ]

    for item in customer_responsibilities_cc:
        add_bullet(doc, item)
    add_blank(doc)

    # ── iv. Project Schedule / Staffing Plan ──────────────────────────────────
    add_para(doc, "iv.  Code Conversion Project Schedule:", bold=True)
    add_blank(doc)

    staffing_plan = data.get("staffing_plan", [])
    if staffing_plan:
        headers = ["Week", "Activity", "Snowflake Hours"]
        rows = [
            [str(r.get("week", "")), str(r.get("activity", "")), str(r.get("sf_hours", ""))]
            for r in staffing_plan
        ]
        add_table(doc, headers, rows, col_widths=[1200, 7400, 2200], tracked_cols={0, 1, 2})
    else:
        add_para(doc, "[Staffing plan with hours by week to be inserted]", tracked=True)
    add_blank(doc)

    # ── v. Code Correction Support ────────────────────────────────────────────
    add_para(doc, "v.  Code Correction Support:", bold=True)
    add_blank(doc)

    weeks = data.get("code_correction_support_weeks", 0)

    if eng == "fixed_fee":
        add_para(doc,
            "As part of the Code Conversion Services and during the Code Correction Support "
            "Period, Snowflake will provide the following code correction support to assist in "
            "resolving a defect or error in converted Objects that: (a) was incorrectly "
            "translated by Snowflake resulting in incorrect data results as compared to the "
            "Legacy System or (b) is the root cause of a performance issue prevalent throughout "
            "the entire code base that could be converted with an alternate translation "
            "(in each case, a \"Defect\") (collectively, \"Code Correction Support\")."
        )
        add_blank(doc)
        weeks_text = str(weeks) if weeks else "[###]"
        para = doc.add_paragraph()
        _apply_para_spacing(para)
        _add_plain_run(para,
            "The \"Code Correction Support Period\" will commence as converted objects are first "
            "delivered to the Customer and will end "
        )
        _add_tracked_run(para, weeks_text, bold=True)
        _add_plain_run(para,
            " weeks from final delivery of converted objects to the Customer. Any Defects must "
            "be submitted to Snowflake during the Code Correction Support Period and must be "
            "accompanied by triage documentation."
        )
    else:
        add_para(doc,
            "As part of the Code Conversion Services and during the Code Correction Support "
            "Period, Snowflake will provide the following code correction support on a time and "
            "material basis to assist in resolving a defect or error in converted Objects that "
            "was incorrectly translated by Snowflake resulting in incorrect data results as "
            "compared to the Legacy System (a \"Defect\") (collectively, \"Code Correction "
            "Support\")."
        )
        add_blank(doc)
        add_para(doc,
            "The \"Code Correction Support Period\" will commence as converted objects are first "
            "delivered to the Customer and at the end of the engagement as indicated in the "
            "project schedule above. Any Defects must be submitted to Snowflake during the Code "
            "Correction Support Period and should be accompanied by triage documentation. Code "
            "Correction Support is subject to the available hours for Code Conversion Services "
            "as specified in the Order Form and does not continue past the Term of this SOW."
        )
    add_blank(doc)

    # ── vi. Exclusions ────────────────────────────────────────────────────────
    add_para(doc, "vi.  Exclusions:", bold=True)
    add_blank(doc)

    if eng == "fixed_fee":
        add_para(doc, "Notwithstanding anything to the contrary, Snowflake is not responsible for:")
    else:
        add_para(doc,
            "The following items are not included within the scope of hours for the Code "
            "Conversion Services, including Code Correction Support. Customer acknowledges that "
            "additional hours may be required if Snowflake is requested to perform work and/or "
            "provide support on these items and that some of these items may not be resolvable:"
        )
    add_blank(doc)

    standard_exclusions_cc = [
        "Any additions, updates or modifications to the Objects after the Object extract has been taken by Snowflake;",
        "Operation, incorporation or use of the converted Objects in or with another technology other than the applicable Target Platform;",
        "Dynamic SQL scenarios where referenced code is not available to Snowflake as part of the provided Legacy Code;",
        "Formatting of Objects \u2013 the conversion process does not guarantee that formatting will be kept upon translation;",
        "Performance-related issues (including query performance and optimization) that do not apply to the entirety of the code base;",
        "Data format differences due to explicit format definitions in the Legacy Platform;",
        "Numeric precision where a Legacy Platform rounding results in less accurate information;",
        "Data issues related to non-enforcement of constraints including but not limited to primary keys, foreign keys, and data checks;",
        "Data which is too large for the correlating Snowflake Service counterpart (Blobs, Clobs, Numeric w/o fixed precision, etc.);",
        "System object references with no Snowflake Service equivalent;",
        "Data which is auto-truncated upon insert into Legacy Platform string columns;",
        "Functional features with no Snowflake Service equivalent or automatable workaround;",
        "Differences in database results, sorting and grouping because of differences in character sets or collation;",
        "Issues resulting from environmental factors such as platform/hardware selection or updates, software tool choice, configuration;",
        "Issues/development related to missing objects or objects that do not execute in the source system (invalid code);",
        "Performance tuning work related to data volumes;",
        "Orchestration and/or execution of converted objects and scripts;",
        "Creation of test scenarios and test scripts;",
        "Execution of test scenarios / scripts;",
        "Migration of data;",
        "Updating scripts with environment specific information such as file paths, variables, etc.",
    ]

    for item in standard_exclusions_cc:
        add_bullet(doc, item)

    for item in data.get("extra_exclusions", []):
        add_bullet(doc, item, tracked=True)
    add_blank(doc)

    # ── vii. Service Type ─────────────────────────────────────────────────────
    add_para(doc, "vii.  Service Type:", bold=True)
    add_blank(doc)

    if eng == "fixed_fee":
        add_para(doc,
            "The Code Conversion Services will be provided on a fixed fee basis. Conversion "
            "activities will not exceed the limits and object counts specified in Table 1 above."
        )
    else:
        add_para(doc,
            "The Code Conversion Services are provided on a time and material basis and are "
            "subject to the availability of hours for the Code Conversion Services as specified "
            "in the Order Form."
        )
    add_blank(doc)

    doc.save(output_path)
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
# CLI Entry Point
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    if len(sys.argv) < 3:
        print("Usage: python generate_project_attachments.py <json_input_path> <output_dir>")
        print("  json_input_path: JSON file with 'dmva_attachment' and/or 'code_conversion_attachment' keys")
        print("  output_dir:      Directory where output .docx files will be written")
        sys.exit(1)

    json_path  = sys.argv[1]
    output_dir = sys.argv[2]

    if not os.path.exists(json_path):
        print(f"Error: Input file not found: {json_path}")
        sys.exit(1)

    os.makedirs(output_dir, exist_ok=True)

    with open(json_path) as f:
        top = json.load(f)

    customer = top.get("customer_name", "Customer").replace(" ", "_")
    today    = date.today().strftime("%Y-%m-%d")
    generated = []

    dmva = top.get("dmva_attachment")
    if dmva:
        out = os.path.join(output_dir, f"{customer}_DMVA_Attachment_{today}.docx")
        generate_dmva_attachment(dmva, out)
        print(f"DMVA attachment generated: {out}")
        generated.append(out)

    cc = top.get("code_conversion_attachment")
    if cc:
        out = os.path.join(output_dir, f"{customer}_CodeConversion_Attachment_{today}.docx")
        generate_code_conv_attachment(cc, out)
        print(f"Code Conversion attachment generated: {out}")
        generated.append(out)

    if not generated:
        print("Warning: No 'dmva_attachment' or 'code_conversion_attachment' keys found in JSON.")
        sys.exit(1)


if __name__ == "__main__":
    main()
