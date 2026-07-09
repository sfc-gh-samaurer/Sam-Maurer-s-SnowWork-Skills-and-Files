#!/usr/bin/env python3
"""
Snowflake PS SOW Document Generator

Generates a properly formatted .docx SOW Attachment 1 matching the
Snowflake PS template specification.

Usage:
    python generate_sow.py <json_input_path> <output_path>

The JSON input must conform to the SOW schema (see SKILL.md for details).
"""

import importlib.util
import json
import sys
import os
from copy import deepcopy
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, Emu, Twips
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ── Formatting Constants (from MERGED template) ──────────────────────────
FONT_NAME = "Arial"
FONT_SIZE = Pt(7.5)  # 95250 EMU
PAGE_WIDTH = 7772400   # 8.5 in
PAGE_HEIGHT = 10058400  # 11 in
MARGIN_LR = 457200     # 0.5 in
MARGIN_TOP = 337185     # ~0.37 in
MARGIN_BOTTOM = 640080  # 0.7 in
TABLE_WIDTH_DXA = 10800
LINE_SPACING = 240  # twips = single spacing (Word standard)


def create_document():
    """Create a blank document with correct page setup."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.left_margin = MARGIN_LR
    section.right_margin = MARGIN_LR
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM

    # Remove default empty paragraph
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    return doc


def set_run_format(run, bold=False):
    """Apply standard font formatting to a run."""
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold
    # Force Arial for East Asian text too
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:eastAsia="{FONT_NAME}" w:cs="{FONT_NAME}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
        rFonts.set(qn('w:cs'), FONT_NAME)


def set_paragraph_spacing(paragraph):
    """Apply standard line spacing to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="0" w:after="60" w:line="{LINE_SPACING}" w:lineRule="auto"/>')
        pPr.append(spacing)
    else:
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '60')
        spacing.set(qn('w:line'), str(LINE_SPACING))
        spacing.set(qn('w:lineRule'), 'auto')


def add_heading(doc, text, level=1):
    """Add a heading paragraph. Level 1 = bold, Level 2 = non-bold."""
    para = doc.add_paragraph()
    set_paragraph_spacing(para)
    run = para.add_run(text)
    set_run_format(run, bold=(level == 1))
    return para


def add_body_text(doc, text, bold=False):
    """Add a body text paragraph."""
    para = doc.add_paragraph()
    set_paragraph_spacing(para)
    run = para.add_run(text)
    set_run_format(run, bold=bold)
    return para


def add_blank_line(doc):
    """Add an empty paragraph as a spacer."""
    para = doc.add_paragraph()
    set_paragraph_spacing(para)
    return para


def add_bullet(doc, text):
    """Add a bullet point using Unicode bullet character."""
    para = doc.add_paragraph()
    set_paragraph_spacing(para)
    run = para.add_run(f"\u2022 {text}")
    set_run_format(run)
    return para


def set_table_borders(table):
    """Apply black single-line borders to a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    # Remove existing borders
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def set_table_width(table, width_dxa=TABLE_WIDTH_DXA):
    """Set table to full page width."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
        tblPr.append(tblW)
    else:
        tblW.set(qn('w:w'), str(width_dxa))
        tblW.set(qn('w:type'), 'dxa')


def format_table_cell(cell, text, bold=False):
    """Format a table cell with standard text."""
    # Clear existing content
    for p in cell.paragraphs:
        if p.text:
            for run in p.runs:
                run.text = ""

    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    set_paragraph_spacing(para)

    if isinstance(text, list):
        # Multi-line cell content (e.g., bullet lists within cells)
        for i, item in enumerate(text):
            if i > 0:
                para = cell.add_paragraph()
                set_paragraph_spacing(para)
            run = para.add_run(str(item))
            set_run_format(run, bold=bold)
    else:
        run = para.add_run(str(text))
        set_run_format(run, bold=bold)


def add_table(doc, headers, rows, header_bold=True, col_widths=None):
    """Add a formatted table with headers and data rows."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_table_borders(table)
    set_table_width(table)

    # Header row
    for ci, header_text in enumerate(headers):
        format_table_cell(table.rows[0].cells[ci], header_text, bold=header_bold)

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            if ci < num_cols:
                format_table_cell(table.rows[ri + 1].cells[ci], cell_text)

    # Column widths
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if ci < len(col_widths):
                    tc = cell._tc
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is None:
                        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                        tc.insert(0, tcPr)
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_widths[ci]}" w:type="dxa"/>')
                        tcPr.append(tcW)
                    else:
                        tcW.set(qn('w:w'), str(col_widths[ci]))

    return table


# ── Section Generators ────────────────────────────────────────────────────

def gen_section_header(doc, number, title):
    """Generate a numbered section header."""
    add_blank_line(doc)
    if number:
        add_heading(doc, f"{number}. {title.upper()}", level=1)
    else:
        add_heading(doc, title.upper(), level=1)


def gen_scope_of_services(doc, data):
    """Section 1: Scope of Services."""
    gen_section_header(doc, None, "SCOPE OF SERVICES")
    add_blank_line(doc)

    scope = data.get("scope_of_services", {})

    # 1.1 Executive Summary
    add_heading(doc, "1.1 Executive Summary", level=2)
    add_blank_line(doc)
    for para_text in scope.get("executive_summary", []):
        add_body_text(doc, para_text)
        add_blank_line(doc)

    # Expected Business Outcomes (optional table or bullets)
    outcomes = scope.get("business_outcomes", [])
    if outcomes:
        add_body_text(doc, "Expected Business Outcomes:", bold=True)
        if isinstance(outcomes[0], dict):
            rows = [[o.get("outcome", ""), o.get("description", "")] for o in outcomes]
            add_table(doc, ["Outcome", "Description"], rows)
        else:
            for o in outcomes:
                add_bullet(doc, o)
        add_blank_line(doc)

    # 1.2 Our Understanding
    understanding = scope.get("our_understanding", {})
    if understanding:
        add_heading(doc, "1.2 Our Understanding", level=2)
        for para_text in understanding.get("paragraphs", []):
            add_body_text(doc, para_text)

        challenges = understanding.get("challenges", [])
        if challenges:
            add_blank_line(doc)
            rows = [[c.get("challenge", ""), c.get("description", "")] for c in challenges]
            add_table(doc, ["Challenge", "Description"], rows)

        solution_paragraphs = understanding.get("solution_paragraphs", [])
        if solution_paragraphs:
            add_blank_line(doc)
            add_body_text(doc, "Snowflake Solution:", bold=True)
            for para_text in solution_paragraphs:
                add_body_text(doc, para_text)

        solution_components = understanding.get("solution_components", [])
        if solution_components:
            rows = [[c.get("component", ""), c.get("description", "")] for c in solution_components]
            add_table(doc, ["Component", "Description"], rows)
        add_blank_line(doc)

    # 1.3 Methodology and Engagement Approach
    methodology = scope.get("methodology", {})
    if methodology:
        add_heading(doc, "1.3 Methodology and Engagement Approach", level=2)
        add_blank_line(doc)
        for para_text in methodology.get("paragraphs", []):
            add_body_text(doc, para_text)
            add_blank_line(doc)

        for phase in methodology.get("phases", []):
            add_body_text(doc, phase.get("title", ""), bold=True)
            for para_text in phase.get("paragraphs", []):
                add_body_text(doc, f"\t{para_text}")
            for item in phase.get("activities", []):
                add_bullet(doc, item)
            add_blank_line(doc)


def gen_milestones(doc, data):
    """Section 2: Outcomes and Acceptance Criteria (milestone table)."""
    gen_section_header(doc, "2", "OUTCOMES AND ACCEPTANCE CRITERIA")
    milestones_data = data.get("milestones", {})
    intro = milestones_data.get("intro", "The following table sets forth the milestones, deliverables, and acceptance criteria for this engagement.")
    add_body_text(doc, intro)
    add_blank_line(doc)

    engagement_type = data.get("engagement_type", "fixed_fee")
    milestones = milestones_data.get("items", [])

    if engagement_type == "fixed_fee":
        headers = ["Milestone", "Payment", "Description", "Deliverable"]
        col_widths = [2200, 1200, 4000, 3400]
        rows = []
        for m in milestones:
            rows.append([
                m.get("name", ""),
                m.get("payment", ""),
                m.get("description", ""),
                m.get("deliverable", "")
            ])
        add_table(doc, headers, rows, col_widths=col_widths)
    else:
        # T&M: milestones without payment column
        headers = ["Milestone", "Description", "Deliverable"]
        col_widths = [2400, 4400, 4000]
        rows = []
        for m in milestones:
            rows.append([
                m.get("name", ""),
                m.get("description", ""),
                m.get("deliverable", "")
            ])
        add_table(doc, headers, rows, col_widths=col_widths)
    add_blank_line(doc)


def gen_acceptance_process(doc, data):
    """Section 3: Acceptance Process."""
    gen_section_header(doc, "3", "ACCEPTANCE PROCESS")

    acceptance = data.get("acceptance_process", {})
    subsections = acceptance.get("subsections", [])

    for sub in subsections:
        number = sub.get("number", "")
        title = sub.get("title", "")
        add_heading(doc, f"{number} {title}", level=2)
        for para_text in sub.get("paragraphs", []):
            add_body_text(doc, para_text)
        add_blank_line(doc)


def gen_key_scope_items(doc, data):
    """Section 4: Key Scope Items."""
    gen_section_header(doc, "4", "KEY SCOPE ITEMS")

    scope_items = data.get("key_scope_items", {})

    # In-scope table
    in_scope = scope_items.get("in_scope", [])
    if in_scope:
        add_heading(doc, "4.1 Deliverables", level=2)
        add_blank_line(doc)
        headers = ["Phase", "Scope Item"]
        rows = [[item.get("phase", ""), item.get("scope_item", "")] for item in in_scope]
        add_table(doc, headers, rows, col_widths=[3600, 7200])
        add_blank_line(doc)

    # Out-of-scope — supports both flat strings and structured objects
    out_scope = scope_items.get("out_of_scope", [])
    if out_scope:
        add_heading(doc, "4.2 Out-of-Scope Activities", level=2)
        add_blank_line(doc)
        if isinstance(out_scope[0], dict):
            # Structured format: numbered subsections with title + description
            headers = ["#", "Item", "Description"]
            rows = [[item.get("number", str(i + 1)), item.get("title", ""), item.get("description", "")] for i, item in enumerate(out_scope)]
            add_table(doc, headers, rows, header_bold=False, col_widths=[1000, 3000, 6800])
        else:
            # Simple string format (backward compatible)
            headers = ["#", "Out-of-Scope Item"]
            rows = [[str(i + 1), item] for i, item in enumerate(out_scope)]
            add_table(doc, headers, rows, header_bold=False, col_widths=[800, 10000])
        add_blank_line(doc)


def gen_raci(doc, data):
    """Section 5: RACI and Work Products."""
    gen_section_header(doc, "5", "RESPONSIBILITY ASSIGNMENT (RACI)")

    raci_data = data.get("raci", {})

    # 5.1 RACI intro
    intro = raci_data.get("intro", "")
    if intro:
        add_heading(doc, "5.1 RACI Matrix", level=2)
        add_body_text(doc, intro)
        # Accountability note (e.g., "Customer is Accountable (A) for every activity")
        accountability_note = raci_data.get("accountability_note", "")
        if accountability_note:
            add_body_text(doc, accountability_note, bold=True)
        add_blank_line(doc)

    # RACI table
    raci_items = raci_data.get("items", [])
    if raci_items:
        headers = ["Activity", "Responsible (R)", "Accountable (A)", "Consulted (C)", "Informed (I)"]
        rows = []
        for item in raci_items:
            # Support phase header rows (bold, spans conceptually)
            if item.get("is_phase_header"):
                rows.append([item.get("activity", ""), "", "", "", ""])
            else:
                rows.append([
                    item.get("activity", ""),
                    item.get("responsible", ""),
                    item.get("accountable", ""),
                    item.get("consulted", ""),
                    item.get("informed", "")
                ])
        add_table(doc, headers, rows, header_bold=False, col_widths=[3600, 1800, 1800, 1800, 1800])
        add_blank_line(doc)

    # 5.2 Work Products
    work_products = raci_data.get("work_products", [])
    if work_products:
        add_heading(doc, "5.3 Work Products", level=2)
        add_blank_line(doc)
        headers = ["Phase", "Work Product", "Primary Owner"]
        rows = [[wp.get("phase", ""), wp.get("work_product", ""), wp.get("owner", "")] for wp in work_products]
        add_table(doc, headers, rows, header_bold=False, col_widths=[2400, 5400, 3000])
        add_blank_line(doc)


def gen_roles(doc, data):
    """Section 6: Roles and Responsibilities."""
    gen_section_header(doc, "6", "ROLES AND RESPONSIBILITIES")

    roles_data = data.get("roles", {})

    # Snowflake roles
    sf_roles = roles_data.get("snowflake", [])
    if sf_roles:
        add_heading(doc, "6.1 Snowflake Team", level=2)
        add_blank_line(doc)
        headers = ["Role", "Responsibilities"]
        rows = [[r.get("role", ""), r.get("responsibilities", "")] for r in sf_roles]
        add_table(doc, headers, rows, header_bold=False, col_widths=[3000, 7800])
        add_blank_line(doc)

    # Customer roles
    cust_roles = roles_data.get("customer", [])
    if cust_roles:
        add_heading(doc, "6.2 Customer Team", level=2)
        add_blank_line(doc)
        headers = ["Role", "Responsibilities"]
        rows = [[r.get("role", ""), r.get("responsibilities", "")] for r in cust_roles]
        add_table(doc, headers, rows, col_widths=[3000, 7800])
        add_blank_line(doc)


def gen_governance(doc, data):
    """Section 7: Project Governance."""
    gen_section_header(doc, "7", "PROJECT GOVERNANCE")

    governance = data.get("governance", {})
    intro = governance.get("intro", "")
    if intro:
        add_body_text(doc, intro)
        add_blank_line(doc)

    # 7.1 Project Alignment (optional subsection)
    alignment = governance.get("alignment", {})
    if alignment:
        add_heading(doc, "7.1 Project Alignment", level=2)
        for para_text in alignment.get("paragraphs", []):
            add_body_text(doc, para_text)
        add_blank_line(doc)

    forums = governance.get("forums", [])
    if forums:
        if alignment:
            add_heading(doc, "7.2 Governance Forums", level=2)
            add_blank_line(doc)
        headers = ["Forum", "Cadence", "Key Participants", "Responsibility", "Materials"]
        rows = [[
            f.get("forum", ""),
            f.get("cadence", ""),
            f.get("participants", ""),
            f.get("responsibility", ""),
            f.get("materials", "")
        ] for f in forums]
        add_table(doc, headers, rows, col_widths=[2000, 1600, 2400, 2400, 2400])
        add_blank_line(doc)


def gen_assumptions(doc, data):
    """Section 8: Assumptions."""
    gen_section_header(doc, "8", "ASSUMPTIONS")

    assumptions = data.get("assumptions", [])
    if assumptions:
        headers = ["#", "Assumption"]
        rows = []
        for i, a in enumerate(assumptions):
            if isinstance(a, dict):
                # Structured format: {"number": "8.1", "assumption": "..."}
                rows.append([a.get("number", str(i + 1)), a.get("assumption", "")])
            else:
                # Simple string format (backward compatible)
                rows.append([str(i + 1), a])
        add_table(doc, headers, rows, col_widths=[800, 10000])
        add_blank_line(doc)


def gen_dependencies(doc, data):
    """Section 9: Dependencies."""
    gen_section_header(doc, "9", "DEPENDENCIES")

    deps = data.get("dependencies", [])
    if deps:
        headers = ["#", "Dependency", "Required By"]
        rows = [[str(i + 1), d.get("dependency", ""), d.get("required_by", "")] for i, d in enumerate(deps)]
        add_table(doc, headers, rows, col_widths=[800, 7000, 3000])
        add_blank_line(doc)


def gen_risks(doc, data):
    """Section 10: Risks."""
    gen_section_header(doc, "10", "RISKS AND MITIGATIONS")

    risks = data.get("risks", [])
    if risks:
        headers = ["#", "Risk", "Impact", "Likelihood", "Mitigation"]
        rows = [[
            str(i + 1),
            r.get("risk", ""),
            r.get("impact", ""),
            r.get("likelihood", ""),
            r.get("mitigation", "")
        ] for i, r in enumerate(risks)]
        add_table(doc, headers, rows, header_bold=False, col_widths=[600, 3000, 1800, 1800, 3600])
        add_blank_line(doc)


def gen_access_security(doc, data):
    """Section 11: Access and Security."""
    gen_section_header(doc, "11", "ACCESS AND SECURITY REQUIREMENTS")

    security = data.get("access_security", {})
    paragraphs = security.get("paragraphs", [])
    for p in paragraphs:
        add_body_text(doc, p)
    if paragraphs:
        add_blank_line(doc)

    items = security.get("items", [])
    for item in items:
        add_bullet(doc, item)
    if items:
        add_blank_line(doc)


def gen_change_management(doc, data):
    """Section 12: Change Management."""
    gen_section_header(doc, "12", "CHANGE MANAGEMENT")

    cm = data.get("change_management", {})
    for para_text in cm.get("paragraphs", []):
        add_body_text(doc, para_text)
        add_blank_line(doc)


def gen_fees(doc, data):
    """Section 13: Fees."""
    gen_section_header(doc, "13", "PROFESSIONAL SERVICES FEES")

    fees = data.get("fees", {})
    engagement_type = data.get("engagement_type", "fixed_fee")

    for para_text in fees.get("paragraphs", []):
        add_body_text(doc, para_text)
        add_blank_line(doc)

    if engagement_type == "fixed_fee":
        # Payment schedule table
        payments = fees.get("payment_schedule", [])
        if payments:
            headers = ["Milestone", "Payment Percentage"]
            rows = [[p.get("milestone", ""), p.get("percentage", "")] for p in payments]
            add_table(doc, headers, rows, col_widths=[5400, 5400])
            add_blank_line(doc)

        total = fees.get("total", "")
        if total:
            add_body_text(doc, f"Total Fixed Fee: {total}", bold=True)
            add_blank_line(doc)
    else:
        # T&M: rate card or phase-based
        rate_card = fees.get("rate_card", [])
        if rate_card:
            headers = ["Role", "Rate"]
            rows = [[r.get("role", ""), r.get("rate", "")] for r in rate_card]
            add_table(doc, headers, rows, col_widths=[5400, 5400])
            add_blank_line(doc)

        phases = fees.get("phases", [])
        if phases:
            headers = ["Phase", "Estimated Hours", "Estimated Cost"]
            rows = [[p.get("phase", ""), p.get("hours", ""), p.get("cost", "")] for p in phases]
            add_table(doc, headers, rows, col_widths=[4000, 3400, 3400])
            add_blank_line(doc)

        total = fees.get("total", "")
        not_to_exceed = fees.get("not_to_exceed", "")
        if total:
            add_body_text(doc, f"Estimated Total: {total}", bold=True)
        if not_to_exceed:
            add_body_text(doc, f"Not-to-Exceed: {not_to_exceed}", bold=True)
        if total or not_to_exceed:
            add_blank_line(doc)


def gen_term(doc, data):
    """Section 14: Term."""
    gen_section_header(doc, "14", "TERM")

    term = data.get("term", {})
    for para_text in term.get("paragraphs", []):
        add_body_text(doc, para_text)
        add_blank_line(doc)


def gen_general_provisions(doc, data):
    """Section 15: General Provisions."""
    gen_section_header(doc, "15", "GENERAL PROVISIONS")

    provisions = data.get("general_provisions", {})
    for para_text in provisions.get("paragraphs", []):
        add_body_text(doc, para_text)
        add_blank_line(doc)


def gen_signatures(doc, data):
    """Section 16: Signatures."""
    gen_section_header(doc, "16", "SIGNATURES")

    sigs = data.get("signatures", {})
    customer_name = data.get("customer_name", "Customer")
    intro = sigs.get("intro", f"IN WITNESS WHEREOF, the parties have executed this SOW as of the date last signed below.")
    add_body_text(doc, intro)
    add_blank_line(doc)

    # Snowflake signature block
    headers_sf = ["", ""]
    rows_sf = [
        ["Signature:", "_________________________________"],
        ["Name:", "_________________________________"],
        ["Title:", "_________________________________"],
        ["Date:", "_________________________________"],
    ]
    add_body_text(doc, "SNOWFLAKE INC.", bold=True)
    add_table(doc, ["Signature:", "_________________________________"], [
        ["Name:", "_________________________________"],
        ["Title:", "_________________________________"],
        ["Date:", "_________________________________"],
    ], header_bold=False, col_widths=[5400, 5400])
    add_blank_line(doc)

    # Customer signature block
    add_body_text(doc, customer_name.upper(), bold=True)
    add_table(doc, ["Signature:", "_________________________________"], [
        ["Name:", "_________________________________"],
        ["Title:", "_________________________________"],
        ["Date:", "_________________________________"],
    ], header_bold=False, col_widths=[5400, 5400])


def _load_attachment_generator():
    """Load generate_project_attachments.py from the same scripts directory."""
    path = os.path.join(os.path.dirname(__file__), "generate_project_attachments.py")
    if not os.path.exists(path):
        return None
    spec = importlib.util.spec_from_file_location("generate_project_attachments", path)
    mod  = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


# ── Multi-Attachment SOW Generator ───────────────────────────────────────────
#
# When the input JSON contains an "attachments" key, the generator switches to
# the template-based multi-attachment path (matching JD Power v5 format).
# The existing simple path is preserved for backwards compatibility.
#
# Template paragraph indexes (rocket_sow_template.docx):
#   2   → body text (normal, line=170 auto)
#   6   → Heading 1
#   10  → bold label (before=0 after=0)
#   12  → main bullet (numId=18, line=170, ind left=720)
#   118 → attachment top-level header
#   122 → attachment sub-section (i., ii.)
#   145 → category header bullet (numId=12, line=240, ind left=720)
#   156 → sub-bullet (numId=13, line=240, ind left=1440)

TEMPLATE_PATH = "/tmp/rocket_sow_template.docx"
_NS_XML = 'http://www.w3.org/XML/1998/namespace'


def _load_template():
    """Load template pPr/rPr clones. Returns (ppr_dict, rpr_dict, doc)."""
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(
            f"SOW template not found at {TEMPLATE_PATH}. "
            "Export it first:\n"
            "  cd ~/.snowflake/cortex/.mcp-servers/google-workspace\n"
            "  ./node export_gdoc.mjs 1n-BAcsVTN0YuK6Ky_uDQvIHKRGjYjRb5zSmcRbli4t4 /tmp/rocket_sow_template.docx"
        )
    tmpl = Document(TEMPLATE_PATH)

    def _ppr(idx):
        el = tmpl.paragraphs[idx]._p.find(qn('w:pPr'))
        return deepcopy(el) if el is not None else None

    def _rpr(idx, run_idx=0):
        runs = tmpl.paragraphs[idx].runs
        if not runs:
            return None
        el = runs[run_idx]._r.find(qn('w:rPr'))
        return deepcopy(el) if el is not None else None

    ppr = {
        "body":       _ppr(2),
        "h1":         _ppr(6),
        "bold_label": _ppr(10),
        "bullet":     _ppr(12),
        "attach_h1":  _ppr(118),
        "attach_h2":  _ppr(122),
        "cat_hdr":    _ppr(145),
        "sub_bullet": _ppr(156),
    }
    rpr = {
        "normal": _rpr(2),
        "bullet": _rpr(12),
    }
    return ppr, rpr


def _rpr_bold(rpr_normal):
    """Clone the normal rPr and set bold on."""
    rpr = deepcopy(rpr_normal)
    b = OxmlElement('w:b')
    rpr.insert(0, b)
    return rpr


def _make_run(text, rpr_el):
    r = OxmlElement('w:r')
    if rpr_el is not None:
        r.append(deepcopy(rpr_el))
    t = OxmlElement('w:t')
    t.text = text
    t.set(f'{{{_NS_XML}}}space', 'preserve')
    r.append(t)
    return r


def _append_p(doc, ppr_el, runs):
    """Build <w:p> and insert before sectPr (critical for correct table ordering)."""
    p = OxmlElement('w:p')
    if ppr_el is not None:
        p.append(deepcopy(ppr_el))
    for text, rpr_el in runs:
        if text:
            p.append(_make_run(text, rpr_el))
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    if sect_pr is not None:
        sect_pr.addprevious(p)
    else:
        body.append(p)
    return p


def _page_break(doc):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    if sect_pr is not None:
        sect_pr.addprevious(p)
    else:
        body.append(p)


def _tmpl_table(doc, ppr, rpr, headers, rows):
    """Full-width table using template formatting."""
    n = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n)
    try:
        t.style = 'TableNormal'
    except KeyError:
        pass
    tbl_el = t._tbl
    tbl_pr = tbl_el.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl_el.insert(0, tbl_pr)
    existing_w = tbl_pr.find(qn('w:tblW'))
    if existing_w is not None:
        tbl_pr.remove(existing_w)
    tbl_w = OxmlElement('w:tblW')
    tbl_w.set(qn('w:w'), '10800')
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_pr.insert(0, tbl_w)
    borders_xml = (
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(parse_xml(borders_xml))

    def _cell_p(cell, text, bold=False):
        for para in list(cell.paragraphs):
            para._element.getparent().remove(para._element)
        p = OxmlElement('w:p')
        p.append(deepcopy(ppr["body"]))
        r_el = _rpr_bold(rpr["normal"]) if bold else deepcopy(rpr["normal"])
        p.append(_make_run(str(text), r_el))
        cell._tc.append(p)

    for ci, h in enumerate(headers):
        _cell_p(t.rows[0].cells[ci], h, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            _cell_p(t.rows[ri + 1].cells[ci], str(val))
    return t


# ── Template-based paragraph shorthand helpers ───────────────────────────────

def _tb(doc, ppr, rpr, text, bold=False):
    """Body paragraph."""
    r = _rpr_bold(rpr["normal"]) if bold else rpr["normal"]
    _append_p(doc, ppr["body"], [(text, r)])


def _blank(doc, ppr, rpr):
    _append_p(doc, ppr["body"], [('', rpr["normal"])])


def _h1(doc, ppr, rpr, text):
    _append_p(doc, ppr["h1"], [(text, rpr["normal"])])


def _bold_label(doc, ppr, rpr, text):
    _append_p(doc, ppr["bold_label"], [(text, _rpr_bold(rpr["normal"]))])


def _bullet(doc, ppr, rpr, text):
    _append_p(doc, ppr["bullet"], [(text, rpr["bullet"])])


def _attach_h1(doc, ppr, rpr, text):
    """Section/attachment header: left-aligned bold text with bottom border line."""
    p = OxmlElement('w:p')
    src = ppr.get("attach_h1")
    new_ppr = deepcopy(src) if src is not None else OxmlElement('w:pPr')
    # Force left alignment (remove any center/right jc)
    for jc in new_ppr.findall(qn('w:jc')):
        new_ppr.remove(jc)
    jc_el = OxmlElement('w:jc')
    jc_el.set(qn('w:val'), 'left')
    new_ppr.append(jc_el)
    # Add bottom border (line under header)
    for pBdr in new_ppr.findall(qn('w:pBdr')):
        new_ppr.remove(pBdr)
    new_ppr.append(parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
        '</w:pBdr>'
    ))
    p.append(new_ppr)
    p.append(_make_run(text, _rpr_bold(rpr["normal"])))
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    if sect_pr is not None:
        sect_pr.addprevious(p)
    else:
        body.append(p)
    return p


def _attach_h2(doc, ppr, rpr, text):
    _append_p(doc, ppr["attach_h2"], [(text, _rpr_bold(rpr["normal"]))])


def _cat_hdr(doc, ppr, rpr, text):
    _append_p(doc, ppr["cat_hdr"], [(text, _rpr_bold(rpr["normal"]))])


def _sub_bullet(doc, ppr, rpr, text):
    _append_p(doc, ppr["sub_bullet"], [(text, rpr["bullet"])])


# ── Main body (preamble + legal sections) ────────────────────────────────────

def _gen_main_body(doc, ppr, rpr, data):
    """Standard SOW preamble and legal sections (before attachments).

    Uses verbatim text from static_content.py for all boilerplate.
    Section headers use _attach_h1() (no auto-numbering) so the letter in the
    title string is the ONLY label — no double-prefixing.
    Roles are NOT rendered here; they belong in attachments or proposals.

    Section map: A B C(Training Funds) D(Payments) E(Scheduling)
                 F(Snowflake Access) G(Additional Terms)
                 H(AVC — optional) I(Subcontractor — optional)
    """
    import sys as _sys
    _sys.path.insert(0, os.path.dirname(__file__))
    import static_content as sc

    customer_name     = data.get("customer_name", "Customer")
    attachments       = data.get("attachments", [])
    production_access = data.get("production_access", {})
    avc_data          = data.get("assumption_validation_checkpoint", {})
    sub_data          = data.get("subcontractor", {})
    source_envs       = data.get("source_environments", [])
    target_envs       = data.get("target_environments", [])
    duration          = data.get("engagement_duration", sc.SECTION_E_DURATION_DEFAULT)
    training_amount   = data.get("training_funds", {}).get("amount", sc.SECTION_C_AMOUNT_DEFAULT)
    training_expiry   = data.get("training_funds", {}).get("expiry", sc.SECTION_C_EXPIRY_DEFAULT)

    # ── Header & Preamble ────────────────────────────────────────────────────
    _attach_h1(doc, ppr, rpr, sc.HEADER)
    _blank(doc, ppr, rpr)
    _tb(doc, ppr, rpr, sc.PREAMBLE_P1)
    _tb(doc, ppr, rpr, sc.PREAMBLE_P2.format(customer_name=customer_name))
    _blank(doc, ppr, rpr)

    # ── A. Description of Technical Services ─────────────────────────────────
    _attach_h1(doc, ppr, rpr, sc.SECTION_A_TITLE)
    _blank(doc, ppr, rpr)
    _tb(doc, ppr, rpr, sc.SECTION_A_BODY)
    _blank(doc, ppr, rpr)

    # ── B. Custom Fixed Fee ───────────────────────────────────────────────────
    _attach_h1(doc, ppr, rpr, sc.SECTION_B_TITLE)
    _blank(doc, ppr, rpr)
    _tb(doc, ppr, rpr, sc.SECTION_B_INTRO)
    _blank(doc, ppr, rpr)
    milestones_att_number = len(attachments) + 2
    milestones_att_title  = f"Attachment {milestones_att_number}: Milestones and Payment Terms"
    for i, att in enumerate(attachments):
        att_title = att.get("title", f"Attachment {i+2}")
        att_desc  = att.get("brief_description", "")
        _bold_label(doc, ppr, rpr, att_title)
        if att_desc:
            _tb(doc, ppr, rpr, att_desc)
        _blank(doc, ppr, rpr)
    _bold_label(doc, ppr, rpr, milestones_att_title)
    _tb(doc, ppr, rpr, "Consolidated milestone payment schedule and payment terms for all Technical Services under this SOW.")
    _blank(doc, ppr, rpr)

    # ── C. Training Funds ─────────────────────────────────────────────────────
    _attach_h1(doc, ppr, rpr, sc.SECTION_C_TITLE)
    _blank(doc, ppr, rpr)
    _tb(doc, ppr, rpr, sc.SECTION_C_BODY.format(
        training_amount=training_amount,
        training_expiry=training_expiry,
    ))
    _blank(doc, ppr, rpr)

    # ── D. Payments and Expenses ──────────────────────────────────────────────
    _attach_h1(doc, ppr, rpr, sc.SECTION_D_TITLE)
    _blank(doc, ppr, rpr)
    _tb(doc, ppr, rpr, sc.SECTION_D_P1)
    _tb(doc, ppr, rpr, sc.SECTION_D_P2)
    _blank(doc, ppr, rpr)

    # ── E. Scheduling and Term ────────────────────────────────────────────────
    _attach_h1(doc, ppr, rpr, sc.SECTION_E_TITLE)
    _blank(doc, ppr, rpr)
    _tb(doc, ppr, rpr, sc.SECTION_E_BODY.format(engagement_duration=duration))
    _blank(doc, ppr, rpr)

    # ── F. Snowflake Access ───────────────────────────────────────────────────
    _attach_h1(doc, ppr, rpr, sc.SECTION_F_TITLE)
    _blank(doc, ppr, rpr)

    _attach_h2(doc, ppr, rpr, sc.F_1_HEADER)
    if source_envs:
        source_list = ", ".join(source_envs)
        target_env  = ", ".join(target_envs) if target_envs else "Snowflake account(s)"
        _tb(doc, ppr, rpr, sc.F_1_BODY_TEMPLATE.format(
            source_list=source_list, target_env=target_env
        ))
    else:
        _tb(doc, ppr, rpr, sc.F_1_BODY_GENERIC)

    _attach_h2(doc, ppr, rpr, sc.F_2_HEADER)
    _tb(doc, ppr, rpr, sc.F_2_BODY)

    _attach_h2(doc, ppr, rpr, sc.F_3_HEADER)
    _tb(doc, ppr, rpr, sc.F_3_BODY)
    if production_access.get("needed"):
        _tb(doc, ppr, rpr, sc.F_3_PROD_NOTE)

    _attach_h2(doc, ppr, rpr, sc.F_4_HEADER)
    _tb(doc, ppr, rpr, sc.F_4_BODY)
    _blank(doc, ppr, rpr)

    # ── G. Additional Terms ───────────────────────────────────────────────────
    _attach_h1(doc, ppr, rpr, sc.SECTION_G_TITLE)
    _blank(doc, ppr, rpr)

    _attach_h2(doc, ppr, rpr, sc.G_1_HEADER)
    _tb(doc, ppr, rpr, sc.G_1_BODY)

    _attach_h2(doc, ppr, rpr, sc.G_2_HEADER)
    _tb(doc, ppr, rpr, sc.G_2_BODY)

    _attach_h2(doc, ppr, rpr, sc.G_3_HEADER)
    _tb(doc, ppr, rpr, sc.G_3_BODY)

    _attach_h2(doc, ppr, rpr, sc.G_4_HEADER)
    _tb(doc, ppr, rpr, sc.G_4_BODY)

    _attach_h2(doc, ppr, rpr, sc.G_5_HEADER)
    if sub_data.get("enabled") and sub_data.get("name"):
        _tb(doc, ppr, rpr, sc.G_5_WITH_PARTNER.format(
            partner_name=sub_data["name"],
            attachment_reference=sub_data.get("attachment_reference", "the applicable Attachment")
        ))
    else:
        _tb(doc, ppr, rpr, sc.G_5_NO_PARTNER)
    _blank(doc, ppr, rpr)

    # ── H. Fixed Fee Engagement Terms (OPTIONAL — AVC) ───────────────────────
    if avc_data.get("enabled"):
        _attach_h1(doc, ppr, rpr, sc.SECTION_H_TITLE)
        _blank(doc, ppr, rpr)
        _attach_h2(doc, ppr, rpr, sc.H_1_HEADER)
        _tb(doc, ppr, rpr, sc.H_AVC_BODY.format(
            trigger_milestone=avc_data.get("trigger_milestone", "the applicable milestone"),
            duration=avc_data.get("duration", "one (1) week")
        ))
        _blank(doc, ppr, rpr)
        _tb(doc, ppr, rpr, sc.H_AVC_GOVERNANCE_LABEL, bold=True)
        _tb(doc, ppr, rpr, sc.H_AVC_GOVERNANCE_BODY)
        _blank(doc, ppr, rpr)

    # ── I. Subcontractor Technical Services (OPTIONAL) ───────────────────────
    if sub_data.get("enabled") and sub_data.get("name"):
        _attach_h1(doc, ppr, rpr, sc.SECTION_I_TITLE)
        _blank(doc, ppr, rpr)
        _tb(doc, ppr, rpr, sc.I_SUBCONTRACTOR_BODY.format(
            partner_name=sub_data["name"],
            partner_role=sub_data.get("role", "Execution Partner"),
            attachment_reference=sub_data.get("attachment_reference", "the applicable Attachment")
        ))
        _blank(doc, ppr, rpr)



# ── Single attachment section ────────────────────────────────────────────────

def _gen_attachment_section(doc, ppr, rpr, att_data, att_number, library_defaults):
    """Generate one full attachment section."""
    from attachment_library import merge_customer_responsibilities, merge_list

    title = att_data.get("title", f"Attachment {att_number}: {library_defaults['title_default']}")
    _page_break(doc)
    _attach_h1(doc, ppr, rpr, title)
    _blank(doc, ppr, rpr)

    # Intro
    scope_intro = att_data.get("scope_intro", library_defaults.get("scope_intro", ""))
    if scope_intro:
        _tb(doc, ppr, rpr, scope_intro)
        _blank(doc, ppr, rpr)

    # i. Scope table
    scope_table = att_data.get("scope_table", [])
    if scope_table:
        _attach_h2(doc, ppr, rpr, "i.  Scope.")
        _tb(doc, ppr, rpr, "The following table outlines the scope of services under this Attachment.")
        _tmpl_table(doc, ppr, rpr, ["Parameter", "Detail"], scope_table)
        _blank(doc, ppr, rpr)

    # ii. Snowflake Responsibilities
    sf_resp = att_data.get("snowflake_responsibilities", [])
    if sf_resp:
        _attach_h2(doc, ppr, rpr, "ii.  Snowflake Responsibilities:")
        _tb(doc, ppr, rpr, "Snowflake will perform the following tasks subject to the terms of this Attachment:")
        for item in sf_resp:
            _bullet(doc, ppr, rpr, item)
        _blank(doc, ppr, rpr)

    # iii. Customer Responsibilities — merge defaults with extras
    default_cust_resp = library_defaults.get("customer_responsibilities", [])
    extras_cust_resp = att_data.get("customer_responsibilities_extra", [])
    # If full override provided, use it directly
    if att_data.get("customer_responsibilities_override"):
        cust_resp_categories = att_data["customer_responsibilities_override"]
    else:
        cust_resp_categories = merge_customer_responsibilities(default_cust_resp, extras_cust_resp)

    if cust_resp_categories:
        _attach_h2(doc, ppr, rpr, "iii.  Customer Responsibilities:")
        _tb(doc, ppr, rpr,
            "Snowflake's performance of the Technical Services under this Attachment is dependent "
            "on Customer's timely performing of the activities listed below.")
        _blank(doc, ppr, rpr)
        for cat in cust_resp_categories:
            category_name = cat.get("category", "")
            items = cat.get("items", [])
            if items:
                _cat_hdr(doc, ppr, rpr, category_name)
                for item in items:
                    _sub_bullet(doc, ppr, rpr, item)
                _blank(doc, ppr, rpr)

    # iv. Technical Scope Exclusions
    default_excl = library_defaults.get("exclusions", [])
    extras_excl = att_data.get("exclusions_extra", [])
    if att_data.get("exclusions_override"):
        exclusions = att_data["exclusions_override"]
    else:
        exclusions = merge_list(default_excl, extras_excl)

    if exclusions:
        _attach_h2(doc, ppr, rpr, "iv.  Technical Scope Exclusions:")
        _tb(doc, ppr, rpr, "The following Technical Services are not in-scope for purposes of this Attachment:")
        for item in exclusions:
            _bullet(doc, ppr, rpr, item)
        _blank(doc, ppr, rpr)

    # v. Technical Scope Assumptions
    default_assump = library_defaults.get("assumptions", [])
    extras_assump = att_data.get("assumptions_extra", [])
    if att_data.get("assumptions_override"):
        assumptions = att_data["assumptions_override"]
    else:
        assumptions = merge_list(default_assump, extras_assump)

    if assumptions:
        _attach_h2(doc, ppr, rpr, "v.  Technical Scope Assumptions:")
        _tb(doc, ppr, rpr,
            "The parties are proceeding with this Attachment under the following assumptions. "
            "If any assumptions are incorrect, each party acknowledges and agrees that a "
            "Change Order may be required.")
        for item in assumptions:
            _bullet(doc, ppr, rpr, item)
        _blank(doc, ppr, rpr)

    # vi. RACI
    raci_items = att_data.get("raci", library_defaults.get("raci_default", []))
    raci_parties = att_data.get("raci_parties", ["Snowflake SD", "Customer"])

    if raci_items:
        _attach_h2(doc, ppr, rpr, "vi.  Responsibility Assignment (RACI):")
        _tb(doc, ppr, rpr, "R – Responsible. A – Accountable. C – Consulted. I – Informed.")
        # Build table columns: Activity + one per party
        headers = ["Activity"] + raci_parties
        rows = []
        for item in raci_items:
            row = [item.get("activity", "")]
            for party_key in ["sf", "customer", "partner", "partner2"]:
                if len(row) < len(headers):
                    row.append(item.get(party_key, "—"))
            rows.append(row[:len(headers)])
        _tmpl_table(doc, ppr, rpr, headers, rows)
        _blank(doc, ppr, rpr)


# ── Milestones table helper (reused in PM attachment) ────────────────────────

def _gen_milestones_table(doc, ppr, rpr, data):
    """Emit milestone table + payment terms subsections (no page break, no header).
    Called from within the Program Management attachment."""
    engagement_type = data.get("engagement_type", "fixed_fee")
    total_fee       = data.get("total_fee", "$[TBD]")
    milestones_data = data.get("milestones", {})
    milestones      = milestones_data.get("items", [])

    if milestones:
        _attach_h2(doc, ppr, rpr, "vii.  Milestone Payment Schedule:")
        _tb(doc, ppr, rpr, (
            "The following table sets forth the consolidated milestone payment schedule for all "
            "Technical Services under this SOW. Payment is due upon Customer written acceptance "
            "of the Key Deliverables for each milestone."
        ))
        _blank(doc, ppr, rpr)
        if engagement_type == "fixed_fee":
            headers = ["#", "Milestone Name", "Key Deliverables", "%", "Amount ($)", "Target"]
            rows = [[
                m.get("num", ""),
                m.get("name", ""),
                m.get("deliverables", ""),
                m.get("pct", ""),
                m.get("amount", total_fee),
                m.get("target", ""),
            ] for m in milestones]
        else:
            headers = ["#", "Milestone Name", "Key Deliverables", "Target"]
            rows = [[
                m.get("num", ""),
                m.get("name", ""),
                m.get("deliverables", ""),
                m.get("target", ""),
            ] for m in milestones]
        _tmpl_table(doc, ppr, rpr, headers, rows)
        _blank(doc, ppr, rpr)

    _attach_h2(doc, ppr, rpr, "viii.  Payment Terms:")
    _tb(doc, ppr, rpr, (
        "All fees are fixed and milestone-gated as described in the Milestone Payment Schedule "
        "above. Payment is due within thirty (30) days of Snowflake\u2019s invoice following "
        "Customer\u2019s written acceptance of the applicable milestone deliverables. Fixed fees "
        "are inclusive of Snowflake Professional Services delivery management and practice "
        "management oversight."
    ))
    _tb(doc, ppr, rpr, (
        f"The total fixed fee for this engagement is {total_fee}, payable across "
        f"{len(milestones)} milestone(s) as described above."
    ))
    _blank(doc, ppr, rpr)


def _gen_gantt_section(doc, ppr, rpr, gantt_image_path):
    """Embed Gantt chart image from a local file path."""
    from docx.shared import Inches
    if not gantt_image_path or not os.path.exists(gantt_image_path):
        return
    _attach_h2(doc, ppr, rpr, "ix.  Project Timeline:")
    _blank(doc, ppr, rpr)
    try:
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(gantt_image_path, width=Inches(6.5))
        body = doc.element.body
        sect_pr = body.find(qn('w:sectPr'))
        para_el = para._element
        body.remove(para_el)
        if sect_pr is not None:
            sect_pr.addprevious(para_el)
        else:
            body.append(para_el)
    except Exception as e:
        _tb(doc, ppr, rpr, f"[Gantt image could not be embedded: {e}]")
    _blank(doc, ppr, rpr)


def _gen_cross_workstream_governance(doc, ppr, rpr, attachments):
    """Cross-workstream governance subsection — only when 2+ core work streams."""
    _attach_h2(doc, ppr, rpr, "x.  Cross-Workstream Governance:")
    _tb(doc, ppr, rpr, (
        "This engagement involves multiple concurrent work streams. The following governance "
        "structure applies across all work streams to ensure coordinated delivery, timely "
        "dependency resolution, and joint milestone gate management."
    ))
    _blank(doc, ppr, rpr)
    _cat_hdr(doc, ppr, rpr, "Wave Sequencing and Dependency Management")
    _sub_bullet(doc, ppr, rpr, "Snowflake SDM maintains the joint dependency map across all work streams and flags cross-workstream blockers in weekly status calls.")
    _sub_bullet(doc, ppr, rpr, "Milestone gates for downstream work streams are dependent on successful completion of upstream gates. Delays in one work stream may require re-sequencing of subsequent milestones.")
    _blank(doc, ppr, rpr)
    _cat_hdr(doc, ppr, rpr, "Joint Milestone Gates")
    _sub_bullet(doc, ppr, rpr, "Each milestone gate requires written sign-off from the designated Customer stakeholder for the applicable work stream.")
    _sub_bullet(doc, ppr, rpr, "Program-level milestone gates (where multiple work streams converge) require sign-off from the Customer Project Sponsor.")
    _blank(doc, ppr, rpr)
    _cat_hdr(doc, ppr, rpr, "Change Management Across Work Streams")
    _sub_bullet(doc, ppr, rpr, "Change Requests that affect scope, timeline, or cost in one work stream are evaluated for impact across all work streams before approval.")
    _sub_bullet(doc, ppr, rpr, "A signed Change Order is required before any out-of-scope work begins in any work stream.")
    _blank(doc, ppr, rpr)


def _gen_pm_attachment(doc, ppr, rpr, att_data, att_number, library_defaults, data):
    """Generate Program Management attachment including milestones, optional Gantt,
    and optional Cross-Workstream Governance (when 2+ core work streams)."""
    from attachment_library import merge_customer_responsibilities, merge_list

    title = att_data.get("title", f"Attachment {att_number}: Program Management & Cross-Workstream Governance")
    _page_break(doc)
    _attach_h1(doc, ppr, rpr, title)
    _blank(doc, ppr, rpr)

    scope_intro = att_data.get("scope_intro", library_defaults.get("scope_intro", ""))
    if scope_intro:
        _tb(doc, ppr, rpr, scope_intro)
        _blank(doc, ppr, rpr)

    # i. Scope table
    scope_table = att_data.get("scope_table", [])
    if scope_table:
        _attach_h2(doc, ppr, rpr, "i.  Scope.")
        _tmpl_table(doc, ppr, rpr, ["Parameter", "Detail"], scope_table)
        _blank(doc, ppr, rpr)

    # ii. Snowflake Responsibilities
    sf_resp = att_data.get("snowflake_responsibilities", [])
    if sf_resp:
        _attach_h2(doc, ppr, rpr, "ii.  Snowflake Responsibilities:")
        for item in sf_resp:
            _bullet(doc, ppr, rpr, item)
        _blank(doc, ppr, rpr)

    # iii. Customer Responsibilities
    default_cust = library_defaults.get("customer_responsibilities", [])
    extras_cust  = att_data.get("customer_responsibilities_extra", [])
    if att_data.get("customer_responsibilities_override"):
        cust_resp = att_data["customer_responsibilities_override"]
    else:
        cust_resp = merge_customer_responsibilities(default_cust, extras_cust)

    if cust_resp:
        _attach_h2(doc, ppr, rpr, "iii.  Customer Responsibilities:")
        _tb(doc, ppr, rpr,
            "Snowflake\u2019s performance of the Program Management services is dependent on "
            "Customer\u2019s timely performing of the activities listed below.")
        _blank(doc, ppr, rpr)
        for cat in cust_resp:
            items = cat.get("items", [])
            if items:
                _cat_hdr(doc, ppr, rpr, cat.get("category", ""))
                for item in items:
                    _sub_bullet(doc, ppr, rpr, item)
                _blank(doc, ppr, rpr)

    # iv. Technical Scope Exclusions
    exclusions = att_data.get("exclusions_override") or merge_list(
        library_defaults.get("exclusions", []), att_data.get("exclusions_extra", []))
    if exclusions:
        _attach_h2(doc, ppr, rpr, "iv.  Technical Scope Exclusions:")
        for item in exclusions:
            _bullet(doc, ppr, rpr, item)
        _blank(doc, ppr, rpr)

    # v. Technical Scope Assumptions
    assumptions = att_data.get("assumptions_override") or merge_list(
        library_defaults.get("assumptions", []), att_data.get("assumptions_extra", []))
    if assumptions:
        _attach_h2(doc, ppr, rpr, "v.  Technical Scope Assumptions:")
        for item in assumptions:
            _bullet(doc, ppr, rpr, item)
        _blank(doc, ppr, rpr)

    # vi. RACI
    raci_items  = att_data.get("raci", library_defaults.get("raci_default", []))
    raci_parties = att_data.get("raci_parties", ["Snowflake SD", "Customer"])
    if raci_items:
        _attach_h2(doc, ppr, rpr, "vi.  Responsibility Assignment (RACI):")
        _tb(doc, ppr, rpr, "R \u2013 Responsible. A \u2013 Accountable. C \u2013 Consulted. I \u2013 Informed.")
        headers = ["Activity"] + raci_parties
        rows = []
        for item in raci_items:
            row = [item.get("activity", "")]
            for key in ["sf", "customer", "partner", "partner2"]:
                if len(row) < len(headers):
                    row.append(item.get(key, "\u2014"))
            rows.append(row[:len(headers)])
        _tmpl_table(doc, ppr, rpr, headers, rows)
        _blank(doc, ppr, rpr)

    # vii–viii. Milestones and Payment Terms (moved here from separate attachment)
    _gen_milestones_table(doc, ppr, rpr, data)

    # ix. Gantt chart (optional)
    gantt_path = att_data.get("gantt_image_path") or data.get("gantt_image_path")
    _gen_gantt_section(doc, ppr, rpr, gantt_path)

    # x. Cross-Workstream Governance (only when 2+ core work streams)
    all_attachments = data.get("attachments", [])
    core_streams = [a for a in all_attachments if a.get("type") != "program_management"]
    if len(core_streams) >= 2:
        _gen_cross_workstream_governance(doc, ppr, rpr, core_streams)


# ── Multi-attachment SOW entry point ─────────────────────────────────────────

def generate_multi_attachment_sow(data, output_path):
    """
    Generate a multi-attachment SOW using the template-based format.
    Called when data contains an 'attachments' key.

    Structure:
      Order Form Exhibit (main body: sections A-I)
      Attachment 1: Program Management (always — auto-added if not in attachments[])
        - includes Milestones, optional Gantt, optional Cross-Workstream Governance
      Attachment 2+: Core work streams
      Signatures
    """
    import sys
    sys.path.insert(0, os.path.dirname(__file__))
    from attachment_library import get_attachment_defaults, merge_list, merge_customer_responsibilities

    ppr, rpr = _load_template()

    doc = Document(TEMPLATE_PATH)
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('p', 'tbl', 'sdt'):
            body.remove(child)

    s = doc.sections[0]
    s.page_width = Emu(7772400)
    s.page_height = Emu(10058400)
    s.left_margin = Emu(457200)
    s.right_margin = Emu(457200)
    s.top_margin = Emu(337185)
    s.bottom_margin = Emu(640080)

    # Main body (Order Form Exhibit sections A-I)
    _gen_main_body(doc, ppr, rpr, data)

    # Separate PM attachments from core work stream attachments
    attachments = data.get("attachments", [])
    pm_atts   = [a for a in attachments if a.get("type") == "program_management"]
    core_atts = [a for a in attachments if a.get("type") != "program_management"]

    # Ensure there is always exactly one PM attachment (auto-create if missing)
    if not pm_atts:
        pm_att = {
            "type": "program_management",
            "title": "Attachment 1: Program Management",
        }
    else:
        pm_att = pm_atts[0]
        pm_att.setdefault("title", "Attachment 1: Program Management")

    # Generate PM attachment (Attachment 1 — includes milestones)
    pm_defaults = get_attachment_defaults("program_management")
    _gen_pm_attachment(doc, ppr, rpr, pm_att, 1, pm_defaults, data)

    # Generate core work stream attachments (Attachment 2, 3, ...)
    for i, att_data in enumerate(core_atts):
        att_number = i + 2
        att_type = att_data.get("type", "generic")
        library_defaults = get_attachment_defaults(att_type)
        _gen_attachment_section(doc, ppr, rpr, att_data, att_number, library_defaults)

    # Signature blocks (after all attachments)
    _page_break(doc)
    _attach_h1(doc, ppr, rpr, "Signatures")
    _blank(doc, ppr, rpr)
    _tb(doc, ppr, rpr,
        "IN WITNESS WHEREOF, the parties have executed this Statement of Work as of the "
        "Effective Date first written above.")
    _blank(doc, ppr, rpr)
    customer_name = data.get("customer_name", "Customer")
    _bold_label(doc, ppr, rpr, "SNOWFLAKE INC.")
    _tmpl_table(doc, ppr, rpr, ["Signature:", "_________________________________"], [
        ["Name:", "_________________________________"],
        ["Title:", "_________________________________"],
        ["Date:", "_________________________________"],
    ])
    _blank(doc, ppr, rpr)
    _bold_label(doc, ppr, rpr, customer_name.upper())
    _tmpl_table(doc, ppr, rpr, ["Signature:", "_________________________________"], [
        ["Name:", "_________________________________"],
        ["Title:", "_________________________________"],
        ["Date:", "_________________________________"],
    ])

    doc.save(output_path)
    return output_path, []


# ── SOW Validation ───────────────────────────────────────────────────────────

def validate_sow_data(data):
    """
    Validate SOW JSON before generation.

    Returns (issues, warnings, checklist) where:
      issues   — blocking problems that MUST be fixed before CLM submission
      warnings — non-blocking items worth reviewing
      checklist — summary snapshot always shown to the PM

    Usage:
        issues, warnings, checklist = validate_sow_data(data)
        for item in checklist:
            print(item)
        if issues:
            print("BLOCKED:", issues)
    """
    issues   = []
    warnings = []

    # ── Required fields ──────────────────────────────────────────────────────
    cname = data.get("customer_name", "")
    if not cname or cname.strip().lower() == "customer":
        issues.append("customer_name: must be set to the actual customer name (not 'Customer')")

    total_fee = str(data.get("total_fee", ""))
    if not total_fee:
        issues.append("total_fee: required — set to dollar amount or '$[TBD]' with note")
    elif "$[TBD]" in total_fee:
        warnings.append("total_fee is '$[TBD]' — fill in before sending to CLM/Legal")

    # ── Training funds ───────────────────────────────────────────────────────
    tf = data.get("training_funds", {})
    if "$[TBD]" in str(tf.get("amount", "$[TBD]")):
        warnings.append("training_funds.amount is '$[TBD]' — confirm amount with deal team")

    # ── Duration ─────────────────────────────────────────────────────────────
    if not data.get("engagement_duration"):
        warnings.append("engagement_duration not set — Section E will default to 'twelve (12) months'")

    # ── Source environments ──────────────────────────────────────────────────
    if not data.get("source_environments"):
        warnings.append("source_environments not set — Section F.1 will use generic access text")

    # ── Milestones ───────────────────────────────────────────────────────────
    milestones = data.get("milestones", {}).get("items", [])
    if not milestones:
        issues.append("milestones.items: no milestones defined — required for fixed-fee SOW")
    elif len(milestones) < 2:
        warnings.append(f"only {len(milestones)} milestone defined — fixed-fee SOWs typically have 2+")
    for m in milestones:
        num = m.get("num", "?")
        if not m.get("target"):
            warnings.append(f"Milestone {num}: no target date/week — add before customer review")
        if "$[TBD]" in str(m.get("amount", "")):
            warnings.append(f"Milestone {num} amount: '$[TBD]' — fill in before CLM")
        if not m.get("deliverables"):
            warnings.append(f"Milestone {num}: deliverables description is empty")

    # ── Per-attachment validation ─────────────────────────────────────────────
    attachments = data.get("attachments", [])
    if not attachments:
        issues.append("attachments: no work streams defined — add at least one core work stream")
    core_atts = [a for a in attachments if a.get("type") != "program_management"]
    if not core_atts:
        issues.append("attachments: no core work stream (non-program_management) defined")
    for att in attachments:
        att_type = att.get("type", "generic")
        title    = att.get("title", att_type)
        if att_type == "program_management":
            continue
        sf_resp = att.get("snowflake_responsibilities", [])
        if not sf_resp:
            issues.append(f"'{title}': Snowflake Responsibilities is empty — required for all work streams")
        if not att.get("brief_description"):
            warnings.append(f"'{title}': brief_description missing — Section B will show no description")
        if not att.get("scope_table"):
            warnings.append(f"'{title}': scope_table is empty — add source/target platform and fee")

    # ── AVC / Subcontractor ──────────────────────────────────────────────────
    avc = data.get("assumption_validation_checkpoint", {})
    if avc.get("enabled") and not avc.get("trigger_milestone"):
        warnings.append("AVC enabled but trigger_milestone not set — specify which milestone triggers it")

    sub = data.get("subcontractor", {})
    if sub.get("enabled") and not sub.get("name"):
        issues.append("subcontractor.enabled is true but name is not set")
    if sub.get("enabled") and not sub.get("attachment_reference"):
        warnings.append("subcontractor.attachment_reference not set — Section I will use generic reference")

    # ── Build review checklist ───────────────────────────────────────────────
    checklist = [
        f"Customer:           {cname or '[NOT SET]'}",
        f"Total fee:          {total_fee or '[NOT SET]'}",
        f"Duration:           {data.get('engagement_duration', '[default: 12 months]')}",
        f"Milestones:         {len(milestones)} defined",
        f"Work streams:       {len(core_atts)} — {', '.join(a.get('type','?') for a in core_atts) or 'none'}",
        f"AVC (Section H):    {'Yes — after ' + avc.get('trigger_milestone','?') if avc.get('enabled') else 'No'}",
        f"Subcontractor (I):  {sub.get('name','None') if sub.get('enabled') else 'None'}",
        f"Gantt image:        {'Yes — ' + str(data.get('gantt_image_path')) if data.get('gantt_image_path') else 'Not provided'}",
        f"Training funds:     {tf.get('amount', '$[TBD]')}",
        f"Source envs:        {', '.join(data.get('source_environments', [])) or '[generic text]'}",
        f"Legal text version: {_get_legal_version()}",
    ]
    return issues, warnings, checklist


def _get_legal_version():
    """Return the legal text version from static_content.py."""
    try:
        import sys as _sys
        _sys.path.insert(0, os.path.dirname(__file__))
        import static_content as sc
        return getattr(sc, "LEGAL_TEXT_VERSION", "unknown")
    except Exception:
        return "unknown"


def print_validation_report(data):
    """Print a validation report to stdout. Called by the skill before confirming generation."""
    issues, warnings, checklist = validate_sow_data(data)
    print("\n" + "=" * 60)
    print("SOW PRE-GENERATION REVIEW")
    print("=" * 60)
    print("\n--- CHECKLIST ---")
    for line in checklist:
        print(f"  {line}")
    if warnings:
        print("\n--- WARNINGS (review before CLM) ---")
        for w in warnings:
            print(f"  ⚠  {w}")
    if issues:
        print("\n--- ISSUES (must fix before generating) ---")
        for issue in issues:
            print(f"  ✗  {issue}")
        print("\nGeneration BLOCKED. Fix issues above and retry.")
    else:
        print("\n✓ No blocking issues. Ready to generate.")
    print("=" * 60 + "\n")
    return issues, warnings, checklist


def generate_sow(data, output_path):
    """Main generation function."""
    # Route to multi-attachment generator when attachments key is present
    if "attachments" in data:
        return generate_multi_attachment_sow(data, output_path)

    doc = create_document()

    # Title
    add_heading(doc, "SOW ATTACHMENT 1", level=1)
    add_blank_line(doc)

    customer_name = data.get("customer_name", "Customer")
    sow_title = data.get("sow_title", f"Statement of Work — {customer_name}")
    add_body_text(doc, sow_title, bold=True)
    add_blank_line(doc)

    # Generate each section based on what's present in data
    section_generators = [
        ("scope_of_services", gen_scope_of_services),
        ("milestones", gen_milestones),
        ("acceptance_process", gen_acceptance_process),
        ("key_scope_items", gen_key_scope_items),
        ("raci", gen_raci),
        ("roles", gen_roles),
        ("governance", gen_governance),
        ("assumptions", gen_assumptions),
        ("dependencies", gen_dependencies),
        ("risks", gen_risks),
        ("access_security", gen_access_security),
        ("change_management", gen_change_management),
        ("fees", gen_fees),
        ("term", gen_term),
        ("general_provisions", gen_general_provisions),
        ("signatures", gen_signatures),
    ]

    for key, gen_func in section_generators:
        if key in data:
            gen_func(doc, data)

    doc.save(output_path)

    # Auto-generate project attachments if keys are present
    attachments = []
    att_mod = _load_attachment_generator()
    if att_mod:
        out_dir    = os.path.dirname(output_path)
        customer   = data.get("customer_name", "Customer").replace(" ", "_")
        today      = date.today().strftime("%Y-%m-%d")

        if "dmva_attachment" in data:
            att_path = os.path.join(out_dir, f"{customer}_DMVA_Attachment_{today}.docx")
            att_mod.generate_dmva_attachment(data["dmva_attachment"], att_path)
            attachments.append(att_path)

        if "code_conversion_attachment" in data:
            att_path = os.path.join(out_dir, f"{customer}_CodeConversion_Attachment_{today}.docx")
            att_mod.generate_code_conv_attachment(data["code_conversion_attachment"], att_path)
            attachments.append(att_path)

    return output_path, attachments


def main():
    if len(sys.argv) < 3:
        print("Usage: python generate_sow.py <json_input_path> <output_path>")
        print("  json_input_path: Path to JSON file with SOW content")
        print("  output_path: Path for the output .docx file")
        sys.exit(1)

    json_path = sys.argv[1]
    output_path = sys.argv[2]

    if not os.path.exists(json_path):
        print(f"Error: Input file not found: {json_path}")
        sys.exit(1)

    with open(json_path, 'r') as f:
        data = json.load(f)

    # Always validate before generating — show checklist and block on issues
    issues, warnings, checklist = print_validation_report(data)
    if issues:
        sys.exit(1)

    result, attachments = generate_sow(data, output_path)
    print(f"SOW generated: {result}")
    for att in attachments:
        print(f"Attachment generated: {att}")


if __name__ == "__main__":
    main()
