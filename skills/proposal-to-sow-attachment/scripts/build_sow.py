#!/usr/bin/env python3
"""
build_sow.py — Render a legal Statement of Work (Attachment) .docx from a JSON spec.

The spec is produced by the agent after reading a proposal (see extract_proposal.py).
Sections auto-number (1., 2., ...). Supported block types are documented in
references/spec_schema.md.

Usage:
    uv run --project <SKILL_DIR> python <SKILL_DIR>/scripts/build_sow.py \
        --spec /abs/path/spec.json --output /abs/path/out.docx

If --output is omitted, the "output" field in the spec is used.
"""
import argparse
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Default brand palette (overridable via spec["theme"]) ---------------
DEFAULTS = {
    "navy": "11567F", "accent": "29B5E8", "body": "333333",
    "grey": "5B5B5B", "hdr_bg": "11567F", "alt_bg": "EAF4FB",
    "line": "C8C8C8", "font": "Arial",
}


def _rgb(hexstr):
    return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))


class SowBuilder:
    def __init__(self, spec):
        self.spec = spec
        theme = {**DEFAULTS, **spec.get("theme", {})}
        self.FONT = theme["font"]
        self.NAVY = _rgb(theme["navy"])
        self.ACCENT = _rgb(theme["accent"])
        self.BODY = _rgb(theme["body"])
        self.GREY = _rgb(theme["grey"])
        self.HDR_BG = theme["hdr_bg"]
        self.ALT_BG = theme["alt_bg"]
        self.LINE = theme["line"]
        self.ACCENT_HEX = theme["accent"]
        self.WHITE = RGBColor(0xFF, 0xFF, 0xFF)
        self.section_no = 0
        self.doc = Document()
        self._base_styles()

    # ---- styling helpers -------------------------------------------------
    def _base_styles(self):
        n = self.doc.styles["Normal"]
        n.font.name = self.FONT
        n.font.size = Pt(10)
        n.font.color.rgb = self.BODY
        n.paragraph_format.space_after = Pt(6)
        n.paragraph_format.line_spacing = 1.08
        for s in self.doc.sections:
            s.top_margin = Inches(0.9)
            s.bottom_margin = Inches(0.9)
            s.left_margin = Inches(1.0)
            s.right_margin = Inches(1.0)

    def _shade(self, cell, hexcolor):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hexcolor)
        tcPr.append(shd)

    def _borders(self, cell):
        tcPr = cell._tc.get_or_add_tcPr()
        b = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), "4")
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), self.LINE)
            b.append(e)
        tcPr.append(b)

    def _vcenter(self, cell):
        tcPr = cell._tc.get_or_add_tcPr()
        va = OxmlElement("w:vAlign")
        va.set(qn("w:val"), "center")
        tcPr.append(va)

    def _cell_text(self, cell, text, bold=False, size=9, white=False):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        for i, ln in enumerate(str(text).split("\n")):
            r = p.add_run(ln if i == 0 else "\n" + ln)
            r.font.name = self.FONT
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.color.rgb = self.WHITE if white else self.BODY

    # ---- block renderers -------------------------------------------------
    def heading(self, text):
        self.section_no += 1
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(f"{self.section_no}.  {text}")
        r.font.name = self.FONT
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = self.NAVY
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), self.ACCENT_HEX)
        pbdr.append(bottom)
        pPr.append(pbdr)

    def subheading(self, num, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        label = f"{num}  {text}" if num else text
        r = p.add_run(label)
        r.font.name = self.FONT
        r.font.size = Pt(10.5)
        r.font.bold = True
        r.font.color.rgb = self.NAVY

    def paragraph(self, text, num=None):
        p = self.doc.add_paragraph()
        if num:
            r0 = p.add_run(f"{num}  ")
            r0.font.name = self.FONT
            r0.font.size = Pt(10)
            r0.font.bold = True
            r0.font.color.rgb = self.NAVY
        r = p.add_run(text)
        r.font.name = self.FONT
        r.font.size = Pt(10)
        r.font.color.rgb = self.BODY

    def bullets(self, items):
        for it in items:
            p = self.doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(it)
            r.font.name = self.FONT
            r.font.size = Pt(10)
            r.font.color.rgb = self.BODY

    def numbered(self, items):
        # items: [{"num": "2.1", "bold": "Title", "text": "desc"}] ("bold" optional)
        for it in items:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            if it.get("num"):
                r0 = p.add_run(f"{it['num']}  ")
                r0.font.name = self.FONT
                r0.font.size = Pt(10)
                r0.font.bold = True
                r0.font.color.rgb = self.NAVY
            if it.get("bold"):
                rb = p.add_run(f"{it['bold']}. ")
                rb.font.name = self.FONT
                rb.font.size = Pt(10)
                rb.font.bold = True
                rb.font.color.rgb = self.BODY
            if it.get("text"):
                r = p.add_run(it["text"])
                r.font.name = self.FONT
                r.font.size = Pt(10)
                r.font.color.rgb = self.BODY

    def table(self, headers, rows, widths=None, header_size=9, body_size=9):
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(headers):
            c = t.rows[0].cells[j]
            self._shade(c, self.HDR_BG)
            self._borders(c)
            self._vcenter(c)
            self._cell_text(c, h, bold=True, size=header_size, white=True)
        for i, row in enumerate(rows):
            cells = t.add_row().cells
            for j, val in enumerate(row):
                c = cells[j]
                if i % 2 == 1:
                    self._shade(c, self.ALT_BG)
                self._borders(c)
                self._vcenter(c)
                self._cell_text(c, val, size=body_size)
        if widths:
            for j, w in enumerate(widths):
                for row in t.rows:
                    row.cells[j].width = Inches(w)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def signature(self, parties):
        sig = self.doc.add_table(rows=1, cols=len(parties))
        sig.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, party in enumerate(parties):
            c = sig.rows[0].cells[j]
            c.text = ""
            r = c.paragraphs[0].add_run(party)
            r.font.name = self.FONT
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = self.NAVY
            for lbl in ["", "Signature: ____________________________",
                        "Name: ____________________________",
                        "Title: ____________________________",
                        "Date: ____________________________"]:
                pp = c.add_paragraph()
                pp.paragraph_format.space_after = Pt(6)
                rr = pp.add_run(lbl)
                rr.font.name = self.FONT
                rr.font.size = Pt(9.5)
                rr.font.color.rgb = self.BODY

    # ---- cover + footer --------------------------------------------------
    def cover(self, cover):
        if cover.get("title"):
            p = self.doc.add_paragraph()
            r = p.add_run(cover["title"])
            r.font.name = self.FONT
            r.font.size = Pt(20)
            r.font.bold = True
            r.font.color.rgb = self.NAVY
        if cover.get("subtitle"):
            p = self.doc.add_paragraph()
            r = p.add_run(cover["subtitle"])
            r.font.name = self.FONT
            r.font.size = Pt(13)
            r.font.bold = True
            r.font.color.rgb = self.ACCENT
        for lbl, val in cover.get("meta", []):
            pp = self.doc.add_paragraph()
            pp.paragraph_format.space_after = Pt(1)
            rl = pp.add_run(lbl + " ")
            rl.font.name = self.FONT
            rl.font.size = Pt(10)
            rl.font.bold = True
            rl.font.color.rgb = self.NAVY
            rv = pp.add_run(val)
            rv.font.name = self.FONT
            rv.font.size = Pt(10)
            rv.font.color.rgb = self.BODY
        if cover.get("note"):
            p = self.doc.add_paragraph()
            r = p.add_run(cover["note"])
            r.font.name = self.FONT
            r.font.size = Pt(8.5)
            r.font.italic = True
            r.font.color.rgb = self.GREY
        self.doc.add_paragraph().paragraph_format.space_after = Pt(8)

    def footer(self, text):
        fp = self.doc.sections[0].footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run(text)
        r.font.name = self.FONT
        r.font.size = Pt(7.5)
        r.font.color.rgb = self.GREY

    # ---- driver ----------------------------------------------------------
    def render_block(self, b):
        t = b.get("type")
        if t == "paragraph":
            self.paragraph(b["text"], num=b.get("num"))
        elif t == "subheading":
            self.subheading(b.get("num"), b["text"])
        elif t == "bullets":
            self.bullets(b["items"])
        elif t == "numbered":
            self.numbered(b["items"])
        elif t == "table":
            self.table(b["headers"], b["rows"], widths=b.get("widths"),
                       header_size=b.get("header_size", 9),
                       body_size=b.get("body_size", 9))
        elif t == "signature":
            self.signature(b["parties"])
        else:
            raise ValueError(f"Unknown block type: {t}")

    def build(self):
        if self.spec.get("cover"):
            self.cover(self.spec["cover"])
        for section in self.spec.get("sections", []):
            self.heading(section["heading"])
            for b in section.get("blocks", []):
                self.render_block(b)
        if self.spec.get("footer"):
            self.footer(self.spec["footer"])
        return self.doc


def main():
    ap = argparse.ArgumentParser(description="Render a SOW attachment .docx from a JSON spec")
    ap.add_argument("--spec", required=True, help="Path to the JSON spec")
    ap.add_argument("--output", help="Output .docx path (overrides spec['output'])")
    args = ap.parse_args()

    with open(args.spec) as f:
        spec = json.load(f)

    out = args.output or spec.get("output")
    if not out:
        raise SystemExit("No output path: provide --output or spec['output']")

    builder = SowBuilder(spec)
    doc = builder.build()
    doc.save(out)
    print(f"Saved: {out}  ({builder.section_no} sections)")


if __name__ == "__main__":
    main()
